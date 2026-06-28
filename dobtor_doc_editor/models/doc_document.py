import base64
import io
import subprocess
import tempfile
import os
import zipfile
from odoo import models, fields, api
from odoo.exceptions import UserError


# ─── python-docx HTML 轉換輔助函式 ────────────────────────────────────────────

def _add_runs_from_node(para, node):
    """將節點的文字（含行內格式）加入 docx paragraph。"""
    if node.text:
        para.add_run(node.text)
    for child in node:
        ctag = (child.tag or '').lower() if isinstance(child.tag, str) else ''
        if ctag in ('strong', 'b'):
            run = para.add_run(child.text_content())
            run.bold = True
        elif ctag in ('em', 'i'):
            run = para.add_run(child.text_content())
            run.italic = True
        elif ctag == 'u':
            run = para.add_run(child.text_content())
            run.underline = True
        elif ctag == 'br':
            para.add_run('\n')
        elif ctag == 'span':
            style_str = child.get('style', '')
            run = para.add_run(child.text_content())
            if 'bold' in style_str:
                run.bold = True
            if 'italic' in style_str:
                run.italic = True
        else:
            text = child.text_content()
            if text:
                para.add_run(text)
        if child.tail:
            para.add_run(child.tail)


def _add_table_to_docx(doc, node):
    """將 HTML table 節點轉換為 python-docx 表格。"""
    rows = node.xpath('.//tr')
    if not rows:
        return
    max_cols = max(
        sum(int(td.get('colspan', 1)) for td in row.xpath('.//td|.//th'))
        for row in rows
    ) or 1
    table = doc.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    for r_idx, row in enumerate(rows):
        cells = row.xpath('.//td|.//th')
        for c_idx, cell in enumerate(cells):
            if c_idx < max_cols:
                try:
                    table.cell(r_idx, c_idx).text = (cell.text_content() or '').strip()
                except Exception:
                    pass


def _html_node_to_docx(doc, node):
    """遞迴將 lxml HTML 節點轉換為 python-docx 結構。"""
    tag = (node.tag or '').lower() if isinstance(node.tag, str) else ''

    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        para = doc.add_heading(level=level)
        _add_runs_from_node(para, node)

    elif tag == 'table':
        _add_table_to_docx(doc, node)

    elif tag == 'hr':
        doc.add_paragraph('─' * 40)

    elif tag in ('ul', 'ol'):
        for child in node:
            if (child.tag or '').lower() == 'li':
                para = doc.add_paragraph(style='List Bullet')
                _add_runs_from_node(para, child)

    elif tag in ('p', 'li', 'blockquote', 'pre'):
        text = (node.text_content() or '').strip()
        if text:
            para = doc.add_paragraph()
            _add_runs_from_node(para, node)

    elif tag in ('div', 'section', 'article', 'main', 'aside',
                 'header', 'footer', 'body', 'span'):
        # 若 div 有直接文字內容且無塊級子節點，作為段落
        block_tags = {'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                      'table', 'ul', 'ol', 'hr', 'div', 'blockquote', 'pre'}
        has_block_children = any(
            (c.tag or '').lower() in block_tags for c in node
        )
        if not has_block_children:
            text = (node.text_content() or '').strip()
            if text:
                para = doc.add_paragraph()
                _add_runs_from_node(para, node)
        else:
            for child in node:
                _html_node_to_docx(doc, child)

    # script/style/meta 等略過

# 超過此大小（bytes）自動存入 ir.attachment
CONTENT_SIZE_THRESHOLD = 500 * 1024  # 500 KB

import logging
_logger = logging.getLogger(__name__)


def _fix_docx_table_widths(docx_bytes, usable_w_mm):
    """修正 LibreOffice 生成的 DOCX 中所有表格寬度（含巢狀），使其符合頁面可用寬度。

    LO 24.2 以內建預設邊距（~5mm）計算表格寬，導致表格寬 200mm 而版心只有 170mm。
    CSS / inline style / !important 全部無效；唯一可靠解法是在 OOXML 層直接修正。

    修改範圍（三層同步）：
    1. <w:tblGrid> / <w:gridCol> — 欄寬網格（等比縮放）
    2. <w:tblW>                  — 表格總寬（設為 target_twips, type=dxa）
    3. <w:tcW>                   — 儲存格寬度（type=dxa 者等比縮放）

    失敗保護：任何異常回傳原始 bytes（優雅降級，確保使用者至少能下載檔案）。
    """
    try:
        from lxml import etree
        import zipfile

        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        def W(tag):
            return f'{{{WNS}}}{tag}'

        # mm → twips（1 inch = 1440 twips；1 inch = 25.4 mm）
        target_twips = round(usable_w_mm * 1440 / 25.4)

        # 只讀取 document.xml，其餘 DOCX 內容不解析（省記憶體）
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zin:
            doc_xml_bytes = zin.read('word/document.xml')

        root = etree.fromstring(doc_xml_bytes)

        # 找出所有 w:tbl（含巢狀），lxml 的 findall 以文件順序回傳（前序）
        all_tbls = root.findall(f'.//{W("tbl")}')

        for tbl in all_tbls:
            # ── Step 1: 讀取並等比縮放 <w:tblGrid> ──────────────────────────
            tbl_grid = tbl.find(W('tblGrid'))
            grid_cols = tbl_grid.findall(W('gridCol')) if tbl_grid is not None else []

            current_total = (
                sum(int(c.get(W('w'), 0)) for c in grid_cols)
                if grid_cols else 0
            )

            # current_total 為 0（空表格）或已在合理範圍（±2%）則跳過
            if current_total <= 0:
                continue
            if abs(current_total - target_twips) / target_twips < 0.02:
                continue

            scale = target_twips / current_total

            # 等比縮放 gridCol，最後一欄補足捨入誤差
            total_assigned = 0
            for i, col in enumerate(grid_cols):
                if i < len(grid_cols) - 1:
                    new_w = max(1, round(int(col.get(W('w'), 0)) * scale))
                else:
                    new_w = max(1, target_twips - total_assigned)
                col.set(W('w'), str(new_w))
                total_assigned += new_w

            # ── Step 2: 修正 <w:tblW> ────────────────────────────────────────
            tbl_pr = tbl.find(W('tblPr'))
            if tbl_pr is not None:
                tbl_w = tbl_pr.find(W('tblW'))
                if tbl_w is None:
                    tbl_w = etree.SubElement(tbl_pr, W('tblW'))
                tbl_w.set(W('w'), str(target_twips))
                tbl_w.set(W('type'), 'dxa')

            # ── Step 3: 等比縮放所有 <w:tcW>（type=dxa）──────────────────────
            for tc in tbl.iter(W('tc')):
                tc_pr = tc.find(W('tcPr'))
                if tc_pr is None:
                    continue
                tc_w = tc_pr.find(W('tcW'))
                if tc_w is None or tc_w.get(W('type')) != 'dxa':
                    continue
                old_w = int(tc_w.get(W('w'), 0))
                if old_w <= 0:
                    continue
                tc_w.set(W('w'), str(max(1, round(old_w * scale))))

        # ── 重新打包 DOCX（只替換 document.xml，其餘保持原樣）────────────────
        modified_xml = etree.tostring(
            root, xml_declaration=True, encoding='UTF-8', standalone=True
        )
        output = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zin:
            with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = (
                        modified_xml
                        if item.filename == 'word/document.xml'
                        else zin.read(item.filename)
                    )
                    zout.writestr(item, data)
        output.seek(0)
        return output.read()

    except Exception as e:
        _logger.warning('[DocEditor] _fix_docx_table_widths 失敗，回傳原始 DOCX: %s', e)
        return docx_bytes  # 優雅降級：確保使用者至少可下載（邊距可能有誤但檔案可開啟）


class DocDocument(models.Model):
    _name = 'doc.document'
    _description = '文件'
    _inherit = ['mail.thread', 'mail.activity.mixin', 'doc.render.mixin']
    _order = 'write_date desc'

    name = fields.Char(string='文件名稱', required=True, default='未命名文件', tracking=True)

    # 內容欄位（sanitize=False，由 Server-Side Sanitizer 在 create/write 中處理）
    content_html = fields.Html(
        string='內容 (HTML 備份)',
        sanitize=False,
        sanitize_attributes=False,
    )
    # Canvas 引擎結構化資料（主要儲存與讀取來源）
    content_json = fields.Text(
        string='Canvas 結構化內容 (JSON)',
        help='Canvas 編輯器的結構化文件資料，為主要儲存與讀取來源',
    )

    # DOCX 模板引擎欄位（後端 docxtpl 填充路線）
    template_docx = fields.Binary(
        string='DOCX 模板原始檔',
        attachment=True,
        help='上傳的 DOCX 模板（+++INS+++語法已自動轉換為 Jinja2），由後端 docxtpl 填充',
    )
    template_filename = fields.Char(string='模板檔名')
    template_variables = fields.Text(
        string='模板變數清單 (JSON)',
        help='上傳時自動偵測的 {{ variable }} 名稱 JSON 陣列',
    )

    header_html = fields.Html(
        string='頁首',
        sanitize=False,
        sanitize_attributes=False,
    )
    footer_html = fields.Html(
        string='頁尾',
        sanitize=False,
        sanitize_attributes=False,
    )

    # 文件設定
    model_id = fields.Many2one(
        'ir.model',
        string='關聯模型',
        ondelete='set null',
        help='用於欄位變數渲染的目標模型',
    )
    # Many2oneReference 的 model_field 必須指向「存模型技術名稱字串」的 Char 欄位
    # （如 'res.partner'），不能直接指向 model_id（Many2one→ir.model），否則
    # web_read 會執行 self.env[ir.model(id,)] 而拋 KeyError，導致後台清單/表單 500
    model_name = fields.Char(
        string='關聯模型技術名稱',
        related='model_id.model',
        store=True,
        index=True,
        help='由 model_id 自動帶入的技術名稱，供 res_id (Many2oneReference) 使用',
    )
    res_id = fields.Many2oneReference(
        model_field='model_name',
        string='關聯記錄 ID',
        index=True,
        help='與本文件雙向關聯的目標記錄 ID（搭配 doc.linked.mixin 使用）',
    )
    field_aliases = fields.Json(
        string='欄位別名對映',
        default=dict,
        help=(
            '中文 token 到 Jinja2 表達式的對映，例如 '
            '{"工程名稱": "object.project_id.name"}。'
            '渲染時 _render_template 會先把文件內 《token》 包裝展開成 '
            '{{ expression }} 再交給 Jinja2。'
        ),
    )
    # ─── 版本管理（W7-8 P1-1）─────────────────────────────────────
    # 設計（修正版）：把版本快照陣列直接存在 doc.document.versions_data (fields.Json)
    #       原本想透過 mail.thread message body 內嵌 base64 JSON，但 body 是 HTML
    #       field 會被 sanitizer 處理（剝除自訂 data-* 屬性），無法 round-trip。
    #       現在把實際 snapshot 內容存在 versions_data，message_post 只記摘要。
    version_number = fields.Integer(
        string='版本號',
        default=0,
        copy=False,
        help='文件版本流水號，每次儲存版本快照時 +1',
    )
    versions_data = fields.Json(
        string='版本快照陣列',
        default=list,
        copy=False,
        help='儲存所有版本快照的 JSON 陣列。每筆含 version_no/created_at/'
             'author_id/author_name/label/content_html/content_json。',
    )
    template_id = fields.Many2one(
        'doc.template',
        string='使用範本',
        ondelete='set null',
    )
    page_format = fields.Selection([
        ('A4', 'A4'),
        ('A3', 'A3'),
        ('A5', 'A5'),
        ('letter', 'Letter'),
        ('legal', 'Legal'),
    ], string='頁面格式', default='A4')
    margin_top = fields.Integer(string='上邊距 (px)', default=96)
    margin_bottom = fields.Integer(string='下邊距 (px)', default=96)
    margin_left = fields.Integer(string='左邊距 (px)', default=96)
    margin_right = fields.Integer(string='右邊距 (px)', default=96)
    has_different_first_page = fields.Boolean(string='首頁不同頁首/頁尾', default=False)
    first_header_html = fields.Html(string='首頁頁首', sanitize=False)
    first_footer_html = fields.Html(string='首頁頁尾', sanitize=False)

    # 協作者
    collaborator_ids = fields.Many2many(
        'res.users',
        'doc_document_collaborator_rel',
        'document_id', 'user_id',
        string='協作者',
    )

    # 多公司隔離
    company_id = fields.Many2one(
        'res.company',
        string='公司',
        required=True,
        default=lambda self: self.env.company,
        index=True,
    )

    # 大文件 attachment 策略（>500KB 自動存入 ir.attachment）
    content_attachment_id = fields.Many2one(
        'ir.attachment',
        string='內容附件（大文件）',
        ondelete='set null',
    )
    is_large_document = fields.Boolean(string='大文件', default=False)

    # 多欄排版預設值
    default_column_count = fields.Selection([
        ('1', '1 欄'), ('2', '2 欄'), ('3', '3 欄'),
    ], string='預設欄數', default='1')
    default_column_gap = fields.Integer(string='欄間距 (px)', default=24)
    column_rule_style = fields.Selection([
        ('none', '無'), ('solid', '實線'), ('dashed', '虛線'), ('dotted', '點線'),
    ], string='欄分隔線', default='none')

    # 需要 sanitize 的 HTML 欄位清單
    _HTML_FIELDS = ('content_html', 'header_html', 'footer_html',
                    'first_header_html', 'first_footer_html')

    # ─── CRUD Hooks ─────────────────────────────────────────────────

    @api.model_create_multi
    def create(self, vals_list):
        """建立文件時：1) sanitize HTML 欄位 2) 若指定 template_id 但 content_html
        為空，自動把 template.content_html / page_format 複製過來。

        為何在 create() 而非 onchange：onchange 只在後台 Form view 觸發，
        Portal Modal / API / 程式內部 .create() 都不會走 onchange；放在 create()
        裡可以涵蓋所有建立路徑（後台、Portal、ChienYi mixin、單元測試）。
        """
        sanitizer = self.env['doc.sanitizer']
        Template = self.env['doc.template']
        for vals in vals_list:
            # 範本自動填充：template_id 有給但 content_html 沒給（或空）
            if vals.get('template_id') and not vals.get('content_html'):
                template = Template.browse(vals['template_id'])
                if template.exists():
                    if template.content_html:
                        vals['content_html'] = template.content_html
                    if template.page_format and not vals.get('page_format'):
                        vals['page_format'] = template.page_format
            for field in self._HTML_FIELDS:
                if vals.get(field):
                    vals[field] = sanitizer.sanitize_html(vals[field])
        return super().create(vals_list)

    @api.onchange('template_id')
    def _onchange_template_id(self):
        """後台 Form view 選 / 切換範本時，直接覆寫 content_html / page_format。

        為何不保護「已有內容」：在 Form view 顯式切換 template_id 是使用者
        刻意動作，意圖就是要新範本內容。若怕誤覆寫已編輯文件，使用者應該
        先儲存版本快照（Ctrl+Shift+S）或不要切 template。

        實際情境（Sprint 16 bug）：使用者建文件 → 選範本 A → 切到範本 B →
        舊邏輯因 content_html 已有 A 的內容而拒絕覆寫 → 編輯器顯示 A 而非 B
        → 使用者誤以為「範本沒套用」。

        例外：當前 content_html 看起來是「使用者編輯過」（不等於任何 template
        的 content_html），仍直接覆寫；保護機制留給「儲存版本快照」工作流。
        """
        for rec in self:
            if not rec.template_id:
                continue
            tpl = rec.template_id
            if tpl.content_html:
                rec.content_html = tpl.content_html
            if tpl.page_format:
                rec.page_format = tpl.page_format

    def write(self, vals):
        sanitizer = self.env['doc.sanitizer']
        for field in self._HTML_FIELDS:
            if vals.get(field):
                vals[field] = sanitizer.sanitize_html(vals[field])

        # 大文件自動存入 attachment
        if 'content_html' in vals and vals['content_html']:
            html = vals['content_html']
            if len(html.encode('utf-8')) > CONTENT_SIZE_THRESHOLD:
                # 各筆記錄分別建立/更新 attachment，不共用 vals
                base_vals = dict(vals)
                base_vals.pop('content_html', None)
                base_vals['is_large_document'] = True
                base_vals['content_html'] = ''
                for rec in self:
                    att = rec._store_content_as_attachment(html)
                    super(DocDocument, rec).write(
                        dict(base_vals, content_attachment_id=att.id)
                    )
                return True
            else:
                # 文件縮小：清除舊 attachment
                vals['is_large_document'] = False
                vals['content_attachment_id'] = False

        return super().write(vals)

    def _store_content_as_attachment(self, html):
        """將大文件內容存入 ir.attachment。"""
        self.ensure_one()
        data = base64.b64encode(html.encode('utf-8'))
        if self.content_attachment_id:
            self.content_attachment_id.write({'datas': data})
            return self.content_attachment_id
        return self.env['ir.attachment'].create({
            'name': f'doc_content_{self.id}.html',
            'type': 'binary',
            'datas': data,
            'res_model': 'doc.document',
            'res_id': self.id,
        })

    def get_content_html(self):
        """統一取得 HTML 內容（自動判斷來源）。"""
        self.ensure_one()
        if self.is_large_document and self.content_attachment_id:
            return base64.b64decode(self.content_attachment_id.datas).decode('utf-8')
        return self.content_html or ''

    # ─── L2-v2 工具方法：alias 自動生成 + 掃描轉換 ─────────────────────
    def init_aliases_from_model(self, overwrite=False):
        """根據 doc.model_id 自動生成 field_aliases 對映。

        參數：
          overwrite=False 時，既有 token 保留不動；新欄位才加入。
          overwrite=True 時，整批以 model 欄位重建（會清掉使用者自訂的 token）。

        生成規則：
          - 一般欄位：「中文 description」→ `object.field_name`
          - date / datetime：用 format_date(object.x)
          - selection：用 selection_label('x')
          - many2one：「中文 description」→ `object.field_name.display_name`
        """
        self.ensure_one()
        if not self.model_id:
            return {'success': False, 'error': '此文件未設定 model_id'}
        model_name = self.model_id.model
        if model_name not in self.env:
            return {'success': False, 'error': f"模型 '{model_name}' 不存在"}

        existing = dict(self.field_aliases or {})
        added = []
        skipped = []
        IrModelFields = self.env['ir.model.fields']
        ttypes = ('char', 'text', 'integer', 'float', 'monetary',
                  'date', 'datetime', 'boolean', 'selection', 'many2one')
        fields = IrModelFields.search([
            ('model', '=', model_name),
            ('store', '=', True),
            ('ttype', 'in', list(ttypes)),
        ], order='field_description asc')

        for f in fields:
            label = (f.field_description or '').strip()
            if not label:
                continue
            # 計算 expression
            if f.ttype == 'date' or f.ttype == 'datetime':
                expr = f'format_date(object.{f.name})'
            elif f.ttype == 'selection':
                expr = f"selection_label('{f.name}')"
            elif f.ttype == 'many2one':
                expr = f'object.{f.name}.display_name'
            else:
                expr = f'object.{f.name}'

            if label in existing and not overwrite:
                skipped.append(label)
                continue
            existing[label] = expr
            added.append(label)

        self.write({'field_aliases': existing})
        return {
            'success': True,
            'aliases': existing,
            'added': added,
            'skipped': skipped,
        }

    def scan_and_convert_to_alias(self):
        """掃 content_html / content_json 內所有 `{{ ... }}` 文字，根據 field_aliases
        反查對應的中文 token，把找得到的整段替換成 《token》。

        反查邏輯：
          - 抓 {{ EXPRESSION }} 內的 EXPRESSION 字串（去前後空白）
          - 在 field_aliases value 中找完全相等的 expression → 取對應 key (token)
          - 沒找到 → 保留原文（不動）

        回傳 {converted: int, skipped: int}
        """
        self.ensure_one()
        import re as _re
        import json as _json
        # 合併 template + 自身 alias map（用 mixin 的合併邏輯）
        aliases = self._collect_field_aliases()
        if not aliases:
            return {'success': False, 'error': '尚無 alias 對映可供反查（請先按「自動生成」或設定範本）'}

        # 拆出 token alias（中文 key）與 varname alias（純識別符 key）
        # 反查策略：
        #   - 直接命中：{{ inner }} 內 inner 是某 alias value（expression）→ 找 token key
        #   - 兩步命中：inner 是 varname_alias key → 取對應 expression → 再從 token_alias 反查
        token_alias = {}  # key=中文 token, value=expression
        varname_alias = {}  # key=純變數名, value=expression
        expr_to_token = {}
        for k, v in aliases.items():
            ks = str(k).strip()
            vs = str(v).strip()
            if not ks or not vs:
                continue
            if _re.match(r'^[A-Za-z_][\w]*$', ks):
                varname_alias[ks] = vs
            else:
                token_alias[ks] = vs
                expr_to_token[vs] = ks

        pat = _re.compile(r'\{\{\s*(.+?)\s*\}\}')
        converted_total = [0]
        skipped_total = [0]

        def _replace_text(text):
            if not text:
                return text
            def _r(m):
                inner = m.group(1).strip()
                # 1) 直接用 expression 反查 token
                token = expr_to_token.get(inner)
                if token:
                    converted_total[0] += 1
                    return f'《{token}》'
                # 2) 變數名兩步反查
                if inner in varname_alias:
                    expr = varname_alias[inner]
                    token = expr_to_token.get(expr)
                    if token:
                        converted_total[0] += 1
                        return f'《{token}》'
                skipped_total[0] += 1
                return m.group(0)
            return pat.sub(_r, text)

        # content_html
        new_html = _replace_text(self.content_html or '')

        # content_json 遞迴
        cj = self.content_json
        if isinstance(cj, str):
            try:
                cj = _json.loads(cj)
            except Exception:
                cj = None

        def _walk(node):
            if isinstance(node, dict):
                for k, v in node.items():
                    if k == 'value' and isinstance(v, str):
                        node[k] = _replace_text(v)
                    else:
                        _walk(v)
            elif isinstance(node, list):
                for item in node:
                    _walk(item)

        if cj:
            _walk(cj)

        vals = {'content_html': new_html}
        if cj:
            vals['content_json'] = cj
        self.write(vals)
        return {
            'success': True,
            'converted': converted_total[0],
            'skipped': skipped_total[0],
        }

    # ─── Actions ─────────────────────────────────────────────────────

    def action_open_editor(self):
        """開啟全螢幕文件編輯器。"""
        self.ensure_one()
        return {
            'type': 'ir.actions.client',
            'tag': 'dobtor_doc_editor.action_doc_editor',
            'context': {
                'doc_id': self.id,
                'doc_name': self.name,
            },
            'target': 'fullscreen',
        }

    def action_quick_preview(self):
        """從 form view 一鍵預覽——開新分頁顯示渲染後 HTML（含 chip 樣式）。

        需要 doc 同時設好 model_id 與 res_id；否則回提示 message。
        """
        self.ensure_one()
        if not self.model_id or not self.res_id:
            from odoo.exceptions import UserError
            raise UserError(
                "此文件未綁定模型或記錄，無法預覽。\n"
                "請先設定「關聯模型」與「關聯記錄 ID」。"
            )
        return {
            'type': 'ir.actions.act_url',
            'url': f'/dobtor_doc/preview/{self.id}',
            'target': 'new',
        }

    def action_export_pdf(self):
        """匯出 PDF 並下載。

        若此文件已綁定 model_id + res_id，自動帶入對應 record，
        讓 alias（《token》/ {{ varname }}）走完整渲染，
        匯出的 PDF 直接呈現實際欄位值，而非 token 原文。
        """
        self.ensure_one()
        record = self._resolve_bound_record()
        pdf_bytes = self._generate_pdf(record=record)
        self.env['doc.editor.export.log'].record_export(
            self, 'pdf', with_alias=record is not None,
            file_size=len(pdf_bytes), source='form_button',
        )
        attachment = self.env['ir.attachment'].create({
            'name': f'{self.name}.pdf',
            'type': 'binary',
            'datas': base64.b64encode(pdf_bytes),
            'res_model': 'doc.document',
            'res_id': self.id,
            'mimetype': 'application/pdf',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'new',
        }

    def action_export_docx(self):
        """匯出 DOCX 並下載。

        若此文件已綁定 model_id + res_id，自動帶入對應 record，
        讓 alias 完整渲染，匯出的 DOCX 直接呈現實際欄位值。
        """
        self.ensure_one()
        record = self._resolve_bound_record()
        docx_bytes = self._generate_docx_via_libreoffice(record=record)
        self.env['doc.editor.export.log'].record_export(
            self, 'docx', with_alias=record is not None,
            file_size=len(docx_bytes), source='form_button',
        )
        attachment = self.env['ir.attachment'].create({
            'name': f'{self.name}.docx',
            'type': 'binary',
            'datas': base64.b64encode(docx_bytes),
            'res_model': 'doc.document',
            'res_id': self.id,
            'mimetype': ('application/vnd.openxmlformats-officedocument'
                         '.wordprocessingml.document'),
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'new',
        }

    def _resolve_bound_record(self):
        """取出此文件綁定的 record（model_id + res_id）；取不到回 None。

        供匯出流程共用：有綁定就讓 alias 走完整渲染輸出實際值，
        沒綁定（或 model/record 已不存在）就回 None，輸出 token 原文。
        """
        self.ensure_one()
        if not (self.model_id and self.res_id):
            return None
        try:
            rec = self.env[self.model_id.model].browse(self.res_id)
            return rec if rec.exists() else None
        except Exception:
            return None

    def action_batch_export_pdf_zip(self):
        """批次匯出：把選取的多份文件各自渲染成 PDF，打包為單一 zip 下載。

        從 list view 的「動作」選單觸發（self 為多筆 recordset）。
        每份文件各自帶入其綁定 record，alias 渲染成實際值。
        單份失敗不中斷整批：以 _ERROR.txt 留下錯誤訊息一併打包。
        """
        if not self:
            raise UserError("請先勾選要匯出的文件。")

        zip_buffer = io.BytesIO()
        used_names = {}
        ok_count = 0
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for doc in self:
                # 檔名去重（同名文件加序號），並清掉路徑分隔字元
                safe_name = (doc.name or f'document_{doc.id}').replace('/', '_').replace('\\', '_')
                count = used_names.get(safe_name, 0)
                used_names[safe_name] = count + 1
                if count:
                    safe_name = f'{safe_name}_{count + 1}'
                try:
                    record = doc._resolve_bound_record()
                    pdf_bytes = doc._generate_pdf(record=record)
                    zf.writestr(f'{safe_name}.pdf', pdf_bytes)
                    self.env['doc.editor.export.log'].record_export(
                        doc, 'pdf', with_alias=record is not None,
                        file_size=len(pdf_bytes), source='batch',
                    )
                    ok_count += 1
                except Exception as exc:
                    zf.writestr(f'{safe_name}_ERROR.txt',
                                f'匯出失敗：{exc}'.encode('utf-8'))

        if not ok_count:
            raise UserError("批次匯出失敗：沒有任何文件成功產生 PDF。")

        zip_bytes = zip_buffer.getvalue()
        filename = (f'documents_{fields.Datetime.now().strftime("%Y%m%d_%H%M%S")}'
                    f'_{ok_count}docs.zip')
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(zip_bytes),
            'mimetype': 'application/zip',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'new',
        }

    # ─── 後端匯出邏輯 ─────────────────────────────────────────────────

    def _build_full_html(self, rendered_body=None):
        """組合完整的 HTML 文件（頁首＋內文＋頁尾），含 CSS 設定。"""
        self.ensure_one()
        body = rendered_body or self.get_content_html()
        header = self.header_html or ''
        footer = self.footer_html or ''

        page_sizes = {
            'A4': ('210mm', '297mm'),
            'A3': ('297mm', '420mm'),
            'A5': ('148mm', '210mm'),
            'letter': ('215.9mm', '279.4mm'),
            'legal': ('215.9mm', '355.6mm'),
        }
        w, h = page_sizes.get(self.page_format, ('210mm', '297mm'))

        return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
@page {{
    size: {w} {h};
}}
body {{
    font-family: 'Microsoft JhengHei', 'Noto Sans TC', 'Arial', sans-serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #333;
    margin: 0;
    padding: 0;
}}
.doc-header {{ margin-bottom: 12px; border-bottom: 1px solid #ddd; padding-bottom: 8px; }}
.doc-footer {{ margin-top: 12px; border-top: 1px solid #ddd; padding-top: 8px; font-size: 10pt; }}
.doc-page-break {{ page-break-after: always; }}
.doc-field-token {{
    background: #e3f2fd;
    border: 1px solid #90caf9;
    border-radius: 3px;
    padding: 1px 4px;
    font-family: monospace;
    font-size: 0.875em;
    color: #1565c0;
}}
table {{ border-collapse: collapse; width: 100%; }}
td, th {{ border: 1px solid #ccc; padding: 6px; }}
</style>
</head>
<body>
{'<div class="doc-header">' + header + '</div>' if header else ''}
<div class="doc-body">{body}</div>
{'<div class="doc-footer">' + footer + '</div>' if footer else ''}
</body>
</html>"""

    def _generate_pdf(self, record=None):
        """使用 Odoo 內建 wkhtmltopdf 產生 PDF。"""
        self.ensure_one()
        body = self.get_content_html()
        if record:
            body = self._render_template(body, record)

        full_html = self._build_full_html(rendered_body=body)
        Report = self.env['ir.actions.report']
        # px → mm（96dpi：1px = 25.4/96 mm）
        px_to_mm = 25.4 / 96
        pdf_bytes = Report._run_wkhtmltopdf(
            [full_html],          # str，不預先 encode（_run_wkhtmltopdf 內部自行 encode）
            landscape=False,
            specific_paperformat_args={
                # 傳入實際 mm 值，由 wkhtmltopdf CLI 套用正確邊距。
                # @page CSS 只保留 size，避免 CSS margin 與 CLI margin 疊加。
                # left/right 需依賴本模組 report_overrides.py 覆寫的支援。
                'data-report-margin-top':    round(self.margin_top    * px_to_mm, 1),
                'data-report-margin-bottom': round(self.margin_bottom * px_to_mm, 1),
                'data-report-margin-left':   round(self.margin_left   * px_to_mm, 1),
                'data-report-margin-right':  round(self.margin_right  * px_to_mm, 1),
            },
        )
        return pdf_bytes

    def _wrap_html_for_libreoffice(self, body_html):
        """為 LibreOffice 建立最小化的標準 HTML 封裝。

        包含 @page CSS 以確保 DOCX 邊距與螢幕上的 doc-page-sheet 一致。
        注意：此 HTML 僅供 LibreOffice 中介轉換，不會存入 DB 或顯示給使用者。
        """
        # px → cm（96dpi：1px = 2.54/96 cm）
        px_to_cm = 2.54 / 96
        page_sizes_cm = {
            'A4':     (21.0,   29.7),
            'A3':     (29.7,   42.0),
            'A5':     (14.8,   21.0),
            'letter': (21.59,  27.94),
            'Letter': (21.59,  27.94),
            'legal':  (21.59,  35.56),
            'Legal':  (21.59,  35.56),
        }
        w_cm, h_cm = page_sizes_cm.get(self.page_format, (21.0, 29.7))
        mt = round(self.margin_top    * px_to_cm, 2)
        mr = round(self.margin_right  * px_to_cm, 2)
        mb = round(self.margin_bottom * px_to_cm, 2)
        ml = round(self.margin_left   * px_to_cm, 2)
        usable_w_cm = round(w_cm - ml - mr, 2)
        return f"""<html xmlns:o="urn:schemas-microsoft-com:office:office"
  xmlns:w="urn:schemas-microsoft-com:office:word"
  xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="utf-8">
<style>
@page {{ size: {w_cm}cm {h_cm}cm; margin: {mt}cm {mr}cm {mb}cm {ml}cm; }}
body {{ font-family: sans-serif; max-width: {usable_w_cm}cm; margin: 0 auto; }}
table {{ border-collapse: collapse; }}
td, th {{ border: 1px solid black; padding: 4px; word-break: break-word; }}
img {{ max-width: 100%; height: auto; }}
</style>
</head>
<body>
{body_html}
</body>
</html>"""

    def _generate_docx_via_libreoffice(self, record=None):
        """使用 LibreOffice headless 將 HTML 轉換為 DOCX。
        LibreOffice 不可用時自動 fallback 至 python-docx。
        """
        self.ensure_one()
        import shutil
        if not shutil.which('soffice'):
            # Fallback：使用 python-docx（已 pip install）
            return self._generate_docx_via_python(record)

        body = self.get_content_html()
        if record:
            body = self._render_template(body, record)

        # 使用專為 LibreOffice 設計的最小化 HTML 封裝
        # （不用 _build_full_html，避免 @page CSS 觸發 LO MIME 誤判）
        lo_html = self._wrap_html_for_libreoffice(body)

        with tempfile.TemporaryDirectory() as tmpdir:
            html_path = os.path.join(tmpdir, 'input.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(lo_html)

            result = subprocess.run(
                [
                    'soffice', '--headless', '--norestore',
                    # .html 副檔名已足夠讓 LO 識別輸入格式（不需 --infilter）
                    # 'MS Word 2007 XML' 是 LO 的 DOCX export filter 正式名稱，
                    # 明確指定可防止 LO 找不到 export filter 而失敗
                    '--convert-to', 'docx:MS Word 2007 XML',
                    '--outdir', tmpdir,
                    html_path,
                ],
                capture_output=True,
                timeout=60,
            )
            if result.returncode != 0:
                raise UserError(
                    f'LibreOffice 轉換失敗：{result.stderr.decode("utf-8", errors="replace")}'
                )

            docx_path = os.path.join(tmpdir, 'input.docx')
            if not os.path.exists(docx_path):
                # LO 有時以原始檔名為基礎產生輸出，嘗試搜尋
                candidates = [
                    f for f in os.listdir(tmpdir)
                    if f.endswith('.docx')
                ]
                if candidates:
                    docx_path = os.path.join(tmpdir, candidates[0])
                else:
                    raise UserError(
                        'LibreOffice 轉換後找不到輸出檔案。'
                        f'（stderr: {result.stderr.decode("utf-8", errors="replace")[:500]}）'
                    )

            with open(docx_path, 'rb') as f:
                raw_bytes = f.read()

        # 後處理：修正 LO 以內建預設邊距計算的表格寬度（200mm → 實際版心寬度）
        px_to_mm = 25.4 / 96
        page_w_map = {
            'A4': 210, 'A3': 297, 'A5': 148,
            'letter': 215.9, 'Letter': 215.9,
            'legal': 215.9, 'Legal': 215.9,
        }
        w_mm = page_w_map.get(self.page_format, 210)
        usable_mm = w_mm - (self.margin_left * px_to_mm) - (self.margin_right * px_to_mm)
        return _fix_docx_table_widths(raw_bytes, usable_mm)

    def _generate_docx_via_python(self, record=None):
        """使用 python-docx + lxml 將 HTML 轉換為 DOCX（LibreOffice 不可用時的 fallback）。"""
        self.ensure_one()
        try:
            from docx import Document
            from docx.shared import Mm
            from lxml import html as lhtml
        except ImportError as e:
            raise UserError(
                f'無法匯出 DOCX：{e}\n'
                '請安裝 python-docx：pip install python-docx'
            )

        body_html = self.get_content_html()
        if record:
            body_html = self._render_template(body_html, record)

        doc = Document()

        # ── 設定頁面大小與邊距 ──
        page_sizes_mm = {
            'A4':     (210, 297),
            'A3':     (297, 420),
            'A5':     (148, 210),
            'letter': (216, 279),
            'legal':  (216, 356),
        }
        w_mm, h_mm = page_sizes_mm.get(self.page_format, (210, 297))
        px_to_mm = 0.264583  # 96dpi：1px = 0.264583mm

        section = doc.sections[0]
        section.page_width   = Mm(w_mm)
        section.page_height  = Mm(h_mm)
        section.top_margin   = Mm(self.margin_top    * px_to_mm)
        section.bottom_margin = Mm(self.margin_bottom * px_to_mm)
        section.left_margin  = Mm(self.margin_left   * px_to_mm)
        section.right_margin = Mm(self.margin_right  * px_to_mm)

        # ── 解析 HTML 並轉換結構 ──
        try:
            tree = lhtml.fromstring(f'<div>{body_html}</div>')
            _html_node_to_docx(doc, tree)
        except Exception:
            # 解析失敗時退回純文字
            from lxml import html as lhtml2
            text = lhtml2.fromstring(f'<div>{body_html}</div>').text_content()
            doc.add_paragraph(text)

        # 若頁首/頁尾有內容，加入文件尾部（簡單實作）
        if self.header_html:
            from lxml import html as lhtml3
            header_text = lhtml3.fromstring(f'<div>{self.header_html}</div>').text_content()
            if header_text.strip():
                para = doc.sections[0].header.paragraphs[0]
                para.text = header_text.strip()
        if self.footer_html:
            from lxml import html as lhtml4
            footer_text = lhtml4.fromstring(f'<div>{self.footer_html}</div>').text_content()
            if footer_text.strip():
                para = doc.sections[0].footer.paragraphs[0]
                para.text = footer_text.strip()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── 版本快照（W7-8 P1-1 重構版）─────────────────────────────────
    # 把實際 snapshot 內容存在 versions_data (fields.Json)，避開 mail.thread
    # body sanitizer 對自訂 data-* 屬性的清洗。mail.thread 仍接收摘要訊息，
    # 提供 audit trail 與 chatter UI 友善體驗。
    #
    # 對外 ID = version_number（自然遞增整數），不再用 message_id 當 key。

    def action_save_version(self, label=None):
        """儲存版本快照。

        Args:
            label: 使用者可自填的版本說明（如「會議結論定稿」）。

        Returns:
            dict: {'version_number': N, 'message_id': msg.id}
        """
        self.ensure_one()
        self.version_number = (self.version_number or 0) + 1
        version_no = self.version_number

        # 寫進 versions_data 陣列（避開 sanitizer）
        # JSON field 會自動序列化；用 list 而非 set 因為要保留順序
        versions = list(self.versions_data or [])
        versions.append({
            'version_no': version_no,
            'created_at': fields.Datetime.now().isoformat(),
            'author_id': self.env.user.id,
            'author_name': self.env.user.name or '匿名',
            'label': label or '',
            'content_html': self.get_content_html() or '',
            'content_json': self.content_json or '',
        })
        self.versions_data = versions

        # 摘要訊息送 chatter（不含 content，只當 audit trail）
        summary = (
            f'<p><strong>版本 v{version_no}</strong>'
            + (f' — {label}' if label else '')
            + '</p>'
        )
        msg = self.message_post(
            body=summary,
            subtype_xmlid='mail.mt_note',
            message_type='comment',
        )
        return {
            'version_number': version_no,
            'message_id': msg.id,
        }

    def get_version_list(self):
        """取得版本快照清單（不含 content，輕量）。"""
        self.ensure_one()
        versions = list(self.versions_data or [])
        # 降序：最新版在前
        versions_sorted = sorted(
            versions,
            key=lambda v: v.get('version_no', 0),
            reverse=True,
        )
        return [
            {
                'version_id': v.get('version_no'),
                'version_number': v.get('version_no'),
                'date': v.get('created_at'),
                'author_id': v.get('author_id'),
                'author_name': v.get('author_name', '匿名'),
                'label': v.get('label', ''),
            }
            for v in versions_sorted
        ]

    def _find_version_entry(self, version_id):
        """內部用：依 version_id 找出對應的 entry（or None）。"""
        for v in (self.versions_data or []):
            if v.get('version_no') == version_id:
                return v
        return None

    def get_version_content(self, version_id):
        """取得單一版本的完整內容（HTML + JSON）。

        Args:
            version_id: 版本號（int）。為向後相容 W7-8 早期 API 也接受傳入字串。
        """
        self.ensure_one()
        try:
            version_id = int(version_id)
        except (TypeError, ValueError):
            return None
        entry = self._find_version_entry(version_id)
        if not entry:
            return None
        return {
            'version_id': entry.get('version_no'),
            'version_number': entry.get('version_no'),
            'date': entry.get('created_at'),
            'author_name': entry.get('author_name', '匿名'),
            'content_html': entry.get('content_html', ''),
            'content_json': entry.get('content_json', ''),
        }

    def restore_version(self, version_id):
        """把指定版本的內容還原到當前文件（會自動先存「還原前」版本）。"""
        self.ensure_one()
        from odoo.exceptions import UserError
        try:
            version_id = int(version_id)
        except (TypeError, ValueError):
            raise UserError("無效的版本識別碼。")
        entry = self._find_version_entry(version_id)
        if not entry:
            raise UserError("找不到指定的版本快照。")

        target_version = entry.get('version_no')
        # 先存「還原前」快照
        self.action_save_version(
            label=f'還原前快照（即將回退到 v{target_version}）'
        )

        # 套用快照內容
        update_vals = {
            'content_html': entry.get('content_html') or '',
            'content_json': entry.get('content_json') or '',
        }
        self.write(update_vals)
        return {
            'restored_version': target_version,
            'new_current_version': self.version_number,
        }

    def diff_versions(self, version_id_a, version_id_b):
        """段落層級 diff。

        Returns:
            dict: {'a_version', 'b_version', 'a_date', 'b_date', 'opcodes': [...]}
                  或 None 若任一 version_id 不存在
        """
        self.ensure_one()
        import difflib
        import re as _re
        from html import unescape

        def _to_lines(html):
            t = _re.sub(
                r'<br\s*/?>|</(p|div|h\d|li|tr)>',
                '\n', html or '', flags=_re.I,
            )
            t = _re.sub(r'<[^>]+>', '', t)
            t = unescape(t)
            return [ln.strip() for ln in t.split('\n') if ln.strip()]

        a = self.get_version_content(version_id_a)
        b = self.get_version_content(version_id_b)
        if not a or not b:
            return None

        lines_a = _to_lines(a.get('content_html'))
        lines_b = _to_lines(b.get('content_html'))

        sm = difflib.SequenceMatcher(None, lines_a, lines_b)
        ops = []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            ops.append({
                'op': tag,
                'a_lines': lines_a[i1:i2],
                'b_lines': lines_b[j1:j2],
            })
        return {
            'a_version': a.get('version_number'),
            'b_version': b.get('version_number'),
            'a_date': a.get('date'),
            'b_date': b.get('date'),
            'opcodes': ops,
        }

    # ─── DB 索引 ─────────────────────────────────────────────────────

    def init(self):
        self.env.cr.execute("""
            CREATE INDEX IF NOT EXISTS doc_document_company_write_date_idx
            ON doc_document (company_id, write_date DESC);
            CREATE INDEX IF NOT EXISTS doc_document_company_model_idx
            ON doc_document (company_id, model_id) WHERE model_id IS NOT NULL;
        """)
