import base64
import io
import json
import logging
import re
import shutil
import subprocess
import tempfile
import os
import zipfile
import html as html_mod
from lxml import etree
from odoo import http
from odoo.http import request

from ..models.doc_zip_guard import (
    assert_input_size,
    inspect_zip_safe,
    ZipBombError,
)

_logger = logging.getLogger(__name__)

# ─── DOCX XML 命名空間 ────────────────────────────────────────────────────────
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML = "http://www.w3.org/XML/1998/namespace"

# DOCX 段落樣式名稱 → HTML 標籤對應表（大小寫不敏感）
_HEADING_STYLES = {
    'heading1': 'h1', 'heading 1': 'h1',
    'heading2': 'h2', 'heading 2': 'h2',
    'heading3': 'h3', 'heading 3': 'h3',
    'heading4': 'h4', 'heading 4': 'h4',
    'heading5': 'h5', 'heading 5': 'h5',
    'heading6': 'h6', 'heading 6': 'h6',
}


def _docx_to_html_with_format(file_bytes):
    """使用 zipfile + lxml 直接解析 DOCX XML，保留字體大小、粗斜體、顏色、對齊等格式。
    無需安裝額外套件：zipfile 為 Python 標準函式庫，lxml 已隨 Odoo 安裝。
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            doc_xml = zf.read('word/document.xml')
            try:
                styles_xml = zf.read('word/styles.xml')
            except KeyError:
                styles_xml = None
    except Exception as e:
        return f'<p>（無法解析 DOCX：{html_mod.escape(str(e))}）</p>'

    # ── 解析段落樣式的預設字體大小 ──────────────────────────────────────────
    style_default_sizes = {}   # styleId → font-size (pt)
    style_display_names = {}   # styleId → displayName (lower)
    if styles_xml:
        try:
            st = etree.fromstring(styles_xml)
            for style in st.findall(f'{{{_W}}}style'):
                sid = style.get(f'{{{_W}}}styleId') or ''
                name_el = style.find(f'{{{_W}}}name')
                display = (name_el.get(f'{{{_W}}}val') or '').lower() if name_el is not None else ''
                style_display_names[sid] = display
                sz = style.find(f'.//{{{_W}}}rPr/{{{_W}}}sz')
                if sz is not None:
                    val = sz.get(f'{{{_W}}}val')
                    if val and val.isdigit():
                        style_default_sizes[sid] = int(val) / 2  # half-point → pt
        except Exception:
            pass

    # ── 解析正文 ─────────────────────────────────────────────────────────────
    try:
        body = etree.fromstring(doc_xml).find(f'{{{_W}}}body')
    except Exception as e:
        return f'<p>（DOCX document.xml 解析失敗：{html_mod.escape(str(e))}）</p>'

    parts = []
    for child in body:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local == 'p':
            parts.append(_w_paragraph_to_html(child, style_default_sizes, style_display_names))
        elif local == 'tbl':
            parts.append(_w_table_to_html(child, style_default_sizes, style_display_names))
        # sectPr（節屬性）及其他 meta 略過

    return '\n'.join(parts) if parts else '<p></p>'


def _w_paragraph_to_html(p, style_sizes, style_names):
    """將 <w:p> 轉換為 HTML 段落或標題。"""
    pPr = p.find(f'{{{_W}}}pPr')
    style_id = ''
    align = ''

    if pPr is not None:
        ps = pPr.find(f'{{{_W}}}pStyle')
        if ps is not None:
            style_id = ps.get(f'{{{_W}}}val') or ''
        jc = pPr.find(f'{{{_W}}}jc')
        if jc is not None:
            jc_val = jc.get(f'{{{_W}}}val', '')
            align = {'center': 'center', 'right': 'right',
                     'both': 'justify', 'distribute': 'justify'}.get(jc_val, '')

    # 決定 HTML 標籤
    display_name = style_names.get(style_id, '').lower()
    tag = _HEADING_STYLES.get(display_name, 'p')

    # 處理 runs
    runs_html = _w_runs_to_html(p, style_sizes.get(style_id))

    content = ''.join(runs_html) or '<br>'
    style_attr = f' style="text-align:{align}"' if align else ''
    return f'<{tag}{style_attr}>{content}</{tag}>'


def _w_runs_to_html(p, default_font_size_pt):
    """將段落內的 <w:r> 轉換為帶 inline style 的 HTML 片段清單。"""
    runs = []
    for child in p:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if local == 'r':
            runs.append(_w_run_to_html(child, default_font_size_pt))
        elif local == 'hyperlink':
            # 超連結：遞迴處理內部 runs
            href = child.get(f'{{{_W}}}anchor') or ''
            inner = ''.join(_w_runs_to_html(child, default_font_size_pt))
            if href:
                runs.append(f'<a href="#{html_mod.escape(href)}">{inner}</a>')
            else:
                runs.append(inner)

    return runs


def _w_run_to_html(r, default_font_size_pt):
    """將單一 <w:r> 轉換為 HTML 字串，保留格式。"""
    rPr = r.find(f'{{{_W}}}rPr')

    # 收集文字（含 <w:t> 與 <w:br>）
    text_parts = []
    for child in r:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local == 't':
            # xml:space="preserve" → 保留空白
            preserve = child.get(f'{{{_XML}}}space') == 'preserve'
            t = child.text or ''
            text_parts.append(html_mod.escape(t) if not preserve else
                               html_mod.escape(t).replace(' ', '&nbsp;'))
        elif local == 'br':
            text_parts.append('<br>')

    if not text_parts:
        return ''

    text = ''.join(text_parts)

    if rPr is None:
        return f'<span style="font-size:{default_font_size_pt}pt">{text}</span>' \
            if default_font_size_pt else text

    # ── 解析格式屬性 ─────────────────────────────────────────────────────────
    css = []

    # 字體大小（half-points → pt）
    sz = rPr.find(f'{{{_W}}}sz')
    if sz is not None:
        val = sz.get(f'{{{_W}}}val')
        if val and val.isdigit():
            css.append(f'font-size: {int(val) // 2}pt')
    elif default_font_size_pt:
        css.append(f'font-size: {default_font_size_pt}pt')

    # 文字顏色（排除 "auto"）
    color = rPr.find(f'{{{_W}}}color')
    if color is not None:
        val = color.get(f'{{{_W}}}val', 'auto')
        if val and val.lower() not in ('auto', 'ffffff'):
            css.append(f'color: #{val}')

    # 字體名稱
    fonts = rPr.find(f'{{{_W}}}rFonts')
    if fonts is not None:
        ff = fonts.get(f'{{{_W}}}ascii') or fonts.get(f'{{{_W}}}eastAsia') or ''
        if ff:
            css.append(f"font-family: '{ff}'")

    # 行內格式（粗體、斜體、底線）
    bold = rPr.find(f'{{{_W}}}b') is not None
    italic = rPr.find(f'{{{_W}}}i') is not None
    underline = rPr.find(f'{{{_W}}}u') is not None

    # 套用格式
    style_attr = f' style="{"; ".join(css)}"' if css else ''
    result = f'<span{style_attr}>{text}</span>' if css else text
    if bold:
        result = f'<strong>{result}</strong>'
    if italic:
        result = f'<em>{result}</em>'
    if underline:
        result = f'<u>{result}</u>'

    return result


def _w_table_to_html(tbl, style_sizes, style_names):
    """將 <w:tbl> 轉換為 HTML 表格（使用 <tbody>，全部 <td> 以相容 Odoo html_editor）。"""
    rows = tbl.findall(f'.//{{{_W}}}tr')
    if not rows:
        return ''
    rows_html = []
    for tr in rows:
        cells = tr.findall(f'{{{_W}}}tc')
        cells_html = []
        for tc in cells:
            colspan = 1
            rowspan = 1
            tcPr = tc.find(f'{{{_W}}}tcPr')
            if tcPr is not None:
                gridSpan = tcPr.find(f'{{{_W}}}gridSpan')
                if gridSpan is not None:
                    val = gridSpan.get(f'{{{_W}}}val')
                    if val and val.isdigit():
                        colspan = int(val)
                vMerge = tcPr.find(f'{{{_W}}}vMerge')
                if vMerge is not None and vMerge.get(f'{{{_W}}}val') == 'restart':
                    rowspan = 2  # 簡化處理，精確 rowspan 需追蹤合併狀態
            span_attrs = ''
            if colspan > 1:
                span_attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                span_attrs += f' rowspan="{rowspan}"'
            # 儲存格內可能有多個段落（<p> 包裝讓 ev.target = p，避免 resize 觸發）
            inner_parts = []
            for p in tc.findall(f'{{{_W}}}p'):
                inner_parts.append(_w_paragraph_to_html(p, style_sizes, style_names))
            inner = ''.join(inner_parts) or '<p>&nbsp;</p>'
            cells_html.append(f'<td{span_attrs}>{inner}</td>')
        rows_html.append(f'<tr>{"".join(cells_html)}</tr>')
    return (
        '<table style="width:100%">'
        '<tbody>'
        + ''.join(rows_html)
        + '</tbody>'
        + '</table>'
    )


def _odt_to_html(file_bytes):
    """使用 odfpy 將 ODT 轉換為 HTML，保留對齊、字體大小、行內格式等樣式。

    odfpy 的節點型態皆為 Element，需透過 node.qname[1]（local tag name）
    來識別 p / h / table 等節點。樣式屬性需用 (namespace_uri, localname) tuple 取得。
    """
    try:
        from odf.opendocument import load as odf_load
        from odf.namespaces import FONS, STYLENS, TEXTNS, XLINKNS
        from odf.element import Node as OdfNode
    except ImportError:
        return '<p>（無法解析 ODT：odfpy 未安裝）</p>'

    import html as html_mod

    doc = odf_load(io.BytesIO(file_bytes))
    parts = []

    def _local(node):
        """取得節點的 local tag name（如 'p', 'h', 'table'）。"""
        if hasattr(node, 'qname') and node.qname:
            return node.qname[1].lower()
        return ''

    # ── 建立 stylename → CSS dict 的對應表 ──
    def _extract_css_from_style(style_node):
        """從 style 節點提取 CSS 屬性字典（段落屬性 + 文字屬性）。"""
        css = {}
        for child in style_node.childNodes:
            if child.nodeType != OdfNode.ELEMENT_NODE:
                continue
            ln = _local(child)
            if ln == 'paragraph-properties':
                align = child.attributes.get((FONS, 'text-align'))
                if align:
                    if align == 'start':
                        align = 'left'
                    elif align == 'end':
                        align = 'right'
                    if align in ('center', 'right', 'justify'):
                        css['text-align'] = align
            elif ln == 'text-properties':
                fs = child.attributes.get((FONS, 'font-size'))
                if fs:
                    css['font-size'] = fs
                fw = child.attributes.get((FONS, 'font-weight'))
                if fw == 'bold':
                    css['font-weight'] = 'bold'
                fi = child.attributes.get((FONS, 'font-style'))
                if fi == 'italic':
                    css['font-style'] = 'italic'
                fc = child.attributes.get((FONS, 'color'))
                if fc and fc != 'transparent':
                    css['color'] = fc
        return css

    # 收集所有樣式（named styles → automatic styles 順序，自動樣式可覆蓋）
    style_map = {}
    for source in [getattr(doc, 'styles', None), doc.automaticstyles]:
        if source is None:
            continue
        for node in source.childNodes:
            if node.nodeType != OdfNode.ELEMENT_NODE:
                continue
            name = node.attributes.get((STYLENS, 'name'))
            if name:
                css = _extract_css_from_style(node)
                # 若有 parent-style-name，先繼承父樣式
                parent = node.attributes.get((STYLENS, 'parent-style-name'))
                if parent and parent in style_map:
                    merged = dict(style_map[parent])
                    merged.update(css)
                    style_map[name] = merged
                elif css:
                    style_map[name] = css

    def _para_style_attr(node):
        """從段落/標題節點取得 style="" 屬性字串（僅段落層級屬性）。"""
        sname = node.attributes.get((TEXTNS, 'style-name'))
        if not sname:
            return ''
        css = style_map.get(sname, {})
        if not css:
            return ''
        return ' style="' + '; '.join(f'{k}: {v}' for k, v in css.items()) + '"'

    def _inline_html(node):
        """遞迴將 ODT inline 節點轉換為 HTML，保留 text:span 的字體大小等樣式。"""
        result = []
        for child in node.childNodes:
            if child.nodeType == OdfNode.TEXT_NODE:
                # 純文字節點：直接 escape
                result.append(html_mod.escape(child.data or ''))
            elif child.nodeType == OdfNode.ELEMENT_NODE:
                if not hasattr(child, 'qname') or not child.qname:
                    continue
                ln = child.qname[1].lower()

                if ln == 'span':
                    # text:span — 取得樣式並遞迴處理內容
                    sname = child.attributes.get((TEXTNS, 'style-name'))
                    inner = _inline_html(child)
                    if sname and sname in style_map:
                        # 行內元素：排除 text-align（屬於段落屬性）
                        css = {k: v for k, v in style_map[sname].items()
                               if k != 'text-align'}
                        css_str = '; '.join(f'{k}: {v}' for k, v in css.items())
                        if css_str:
                            result.append(f'<span style="{css_str}">{inner}</span>')
                        else:
                            result.append(inner)
                    else:
                        result.append(inner)
                elif ln == 'line-break':
                    result.append('<br>')
                elif ln == 's':
                    # text:s — 多個空格
                    count_str = child.getAttribute('c') or '1'
                    try:
                        count = max(1, int(count_str))
                    except (ValueError, TypeError):
                        count = 1
                    result.append('&nbsp;' * count)
                elif ln == 'tab':
                    result.append('&nbsp;&nbsp;&nbsp;&nbsp;')
                elif ln == 'a':
                    # text:a — 超連結
                    inner = _inline_html(child)
                    href = child.attributes.get((XLINKNS, 'href'), '#')
                    result.append(
                        f'<a href="{html_mod.escape(href)}">{inner}</a>'
                    )
                else:
                    # 其他行內節點（text:ruby 等）→ 遞迴取文字
                    result.append(_inline_html(child))
        return ''.join(result)

    def _cell_paragraphs(cell_node):
        """將 ODT table cell 內的段落轉換為帶樣式的 HTML（保留 font-size 等）。"""
        result = []
        for child in cell_node.childNodes:
            if child.nodeType != OdfNode.ELEMENT_NODE:
                continue
            ln = _local(child)
            if ln == 'p':
                sa = _para_style_attr(child)
                inline = _inline_html(child)
                result.append(f'<p{sa}>{inline or "<br/>"}</p>')
            elif ln == 'h':
                level = child.getAttribute('outlinelevel') or '1'
                sa = _para_style_attr(child)
                inline = _inline_html(child)
                result.append(f'<h{level}{sa}>{inline or "<br/>"}</h{level}>')
        return ''.join(result) if result else '<p>&nbsp;</p>'

    def _process_table(node):
        """將 ODT table 轉換為 HTML table（使用 <tbody>，全部 <td>，保留段落樣式）。"""
        rows_html = []
        for child in node.childNodes:
            ln = _local(child)
            if ln == 'table-header-rows':
                for row in child.childNodes:
                    if _local(row) == 'table-row':
                        cells = []
                        for c in row.childNodes:
                            if _local(c) == 'table-cell':
                                cells.append(f'<td>{_cell_paragraphs(c)}</td>')
                        if cells:
                            rows_html.append(f'<tr>{"".join(cells)}</tr>')
            elif ln == 'table-row':
                cells = []
                for c in child.childNodes:
                    if _local(c) == 'table-cell':
                        cells.append(f'<td>{_cell_paragraphs(c)}</td>')
                if cells:
                    rows_html.append(f'<tr>{"".join(cells)}</tr>')
        if rows_html:
            parts.append(
                '<table style="width:100%">'
                '<tbody>'
                + ''.join(rows_html)
                + '</tbody>'
                + '</table>'
            )

    def _process(nodes):
        for node in nodes:
            if not hasattr(node, 'qname'):
                continue
            ln = _local(node)
            sa = _para_style_attr(node)
            if ln == 'h':
                level = node.getAttribute('outlinelevel') or '1'
                inline = _inline_html(node)
                if inline:
                    parts.append(f'<h{level}{sa}>{inline}</h{level}>')
            elif ln == 'p':
                inline = _inline_html(node)
                parts.append(f'<p{sa}>{inline}</p>' if inline else '<p><br/></p>')
            elif ln == 'list':
                items = []
                for item in node.childNodes:
                    if _local(item) == 'list-item':
                        for sub in item.childNodes:
                            sln = _local(sub)
                            if sln in ('p', 'h'):
                                t = _inline_html(sub)
                                if t:
                                    items.append(f'<li>{t}</li>')
                if items:
                    parts.append(f'<ul>{"".join(items)}</ul>')
            elif ln == 'table':
                _process_table(node)
            else:
                # 容器節點（text, section, frame 等）→ 繼續往下
                if hasattr(node, 'childNodes') and node.childNodes:
                    _process(node.childNodes)

    _process(doc.body.childNodes)
    return '\n'.join(parts) if parts else '<p></p>'


def _ts_parse_docx_to_elements(file_bytes, float_textbox=False, anchored_image=False):
    """使用本模組的 TS OOXML Parser 把 .docx 解析為 canvas-editor IElement[] JSON。

    流程（Phase E 並行通道）：
        1. 把 file_bytes 寫到暫存檔
        2. subprocess 呼叫 `node tools/dist/parse_docx_cli.cjs <input> <output>`
        3. 讀回 IElement[] JSON 並 parse 為 Python list
        4. 失敗時回 None（caller 應 fallback 到 LibreOffice 路徑）

    依賴：
        - container 內有 Node 18+（`docker exec odoo18 which node` 已驗）
        - `tools/dist/parse_docx_cli.cjs` 已 build（`npm run build:cli` 產出）

    參數：
        file_bytes:       docx 檔案 bytes
        float_textbox:    Sprint Y58 opt-in：展平 wp:anchor + w:txbxContent 文字
        anchored_image:   Sprint Y58 opt-in：透傳 wp:anchor 屬性到 IElement.anchor

    回傳：list[dict] 或 None
    """
    import logging
    _logger = logging.getLogger(__name__)

    # 解析 module 根目錄（doc_controller.py 在 controllers/ 下）
    module_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    cli_path = os.path.join(module_dir, 'tools', 'dist', 'parse_docx_cli.cjs')

    if not os.path.isfile(cli_path):
        _logger.warning(
            'dobtor_doc_editor: TS CLI not built at %s — '
            'run `cd %s && npm run build:cli`',
            cli_path, module_dir,
        )
        return None

    if not shutil.which('node'):
        _logger.warning('dobtor_doc_editor: `node` not in PATH; TS engine unavailable')
        return None

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            in_path = os.path.join(tmpdir, 'input.docx')
            out_path = os.path.join(tmpdir, 'output.json')
            with open(in_path, 'wb') as fp:
                fp.write(file_bytes)

            # Sprint 358-359：--svg-graphics 讓 SmartArt/Chart 渲成 SVG image
            # （取代線性文字 fallback；2026-05-29 真實資料 fidelity audit 修法）。
            # Sprint Y58：--float-textbox / --anchored-image 由 caller 決定是否啟用。
            # 舊版 CLI 遇未知旗標會優雅忽略、不 crash，故部署落差安全。
            argv = ['node', cli_path, in_path, out_path, '--elements', '--svg-graphics']
            if float_textbox:
                argv.append('--float-textbox')
            if anchored_image:
                argv.append('--anchored-image')
            proc = subprocess.run(
                argv,
                capture_output=True,
                timeout=30,
            )
            if proc.returncode != 0:
                _logger.warning(
                    'dobtor_doc_editor: TS CLI failed (rc=%d): %s',
                    proc.returncode,
                    proc.stderr.decode('utf-8', errors='replace')[:500],
                )
                return None

            with open(out_path, 'r', encoding='utf-8') as fp:
                return json.load(fp)
    except subprocess.TimeoutExpired:
        _logger.warning('dobtor_doc_editor: TS CLI timeout (30s)')
        return None
    except Exception as e:
        _logger.warning('dobtor_doc_editor: TS CLI error — %s', e)
        return None


def _lo_convert_to_html(file_bytes, ext):
    """使用 LibreOffice headless 將 ODT/DOCX 轉換為 HTML。

    若 LibreOffice 未安裝，回傳 None（由呼叫端 fallback 至 Python 解析器）。
    轉換後呼叫 _lo_postprocess() 修正表格寬度、字體、頁尾、空白頁等問題。
    """
    if not shutil.which('soffice'):
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, f'input{ext}')
        with open(in_path, 'wb') as fp:
            fp.write(file_bytes)

        proc = subprocess.run(
            ['soffice', '--headless', '--norestore', '--nofirststartwizard',
             '--nolockcheck',
             '--convert-to', 'html:HTML', '--outdir', tmpdir, in_path],
            capture_output=True, timeout=120,
            env={**os.environ, 'HOME': tmpdir},   # 每次獨立 config 目錄，避免鎖定衝突
        )
        if proc.returncode != 0:
            raise Exception(
                f'LibreOffice 轉換失敗：'
                f'{proc.stderr.decode("utf-8", errors="replace")[:400]}'
            )
        html_path = os.path.join(tmpdir, 'input.html')
        if not os.path.exists(html_path):
            raise Exception('LibreOffice：找不到輸出 HTML 檔案')

        with open(html_path, 'r', encoding='utf-8', errors='replace') as fp:
            raw_html = fp.read()

    style_m = re.search(r'<style[^>]*>(.*?)</style>', raw_html, re.DOTALL | re.IGNORECASE)
    margins = _extract_page_margins(style_m.group(1)) if style_m else None
    return _lo_postprocess(raw_html), margins


def _extract_page_margins(style_text):
    """從 @page CSS 規則提取頁面邊距並轉換為 px（96 dpi）。
    支援 cm、mm、in、pt、px 單位。回傳 {'top':N,'right':N,'bottom':N,'left':N} 或 None。
    """
    page_m = re.search(r'@page\s*\{([^}]+)\}', style_text, re.IGNORECASE)
    if not page_m:
        return None

    def _to_px(val):
        val = val.strip()
        try:
            if val.endswith('cm'):  return round(float(val[:-2]) * 37.795)
            if val.endswith('mm'):  return round(float(val[:-2]) * 3.7795)
            if val.endswith('in'):  return round(float(val[:-2]) * 96)
            if val.endswith('pt'):  return round(float(val[:-2]) * 96 / 72)
            if val.endswith('px'):  return round(float(val[:-2]))
        except Exception:
            pass
        return None

    props = {}
    for decl in page_m.group(1).split(';'):
        if ':' not in decl:
            continue
        prop, _, val = decl.partition(':')
        props[prop.strip().lower()] = val.strip()

    margins = {}
    if 'margin' in props:
        parts = props['margin'].split()
        if len(parts) == 1:
            v = _to_px(parts[0])
            if v:
                margins = {'top': v, 'right': v, 'bottom': v, 'left': v}
        elif len(parts) == 2:
            v, h = _to_px(parts[0]), _to_px(parts[1])
            if v and h:
                margins = {'top': v, 'right': h, 'bottom': v, 'left': h}
        elif len(parts) == 4:
            vals = [_to_px(p) for p in parts]
            if all(vals):
                margins = {'top': vals[0], 'right': vals[1],
                           'bottom': vals[2], 'left': vals[3]}
    for side in ('top', 'right', 'bottom', 'left'):
        key = f'margin-{side}'
        if key in props:
            v = _to_px(props[key])
            if v:
                margins[side] = v

    return margins if margins else None


def _lo_postprocess(raw_html):
    """後處理 LibreOffice HTML 輸出，修正已知的格式問題：

    1. <font> → <span>：保留 style 與 color，避免 sanitizer 濾掉後喪失格式
    2. 表格寬度：移除固定 px 寬度，改為 width:100%（避免超出頁面邊界）
    3. page-break-before: always 首段：移除（避免在頂部產生空白頁）
    4. <div title="footer"> 整塊移除（文件原生頁尾不進編輯器）
    5. <sdfield>：移除標籤但保留文字（頁碼欄位）
    6. 巢狀 <ol><ol>：展平為單層（LibreOffice 用多層巢狀表示縮排）
    7. class → CSS inline（針對 body/p/span 等 class 對應 style block）
    8. 回傳 <body> 內容字串
    """
    from io import StringIO

    # 取得 <style> 與 <body>
    style_m = re.search(r'<style[^>]*>(.*?)</style>', raw_html, re.DOTALL | re.IGNORECASE)
    body_m  = re.search(r'<body[^>]*>(.*?)</body>',  raw_html, re.DOTALL | re.IGNORECASE)
    if not body_m:
        return raw_html
    body_str = body_m.group(1).strip()

    # ── CSS class map ────────────────────────────────────────────────────────
    css_map = {}
    if style_m:
        for m in re.finditer(r'([a-zA-Z0-9]*)\.([\w-]+)\s*\{([^}]+)\}',
                             style_m.group(1)):
            tag  = m.group(1).lower() or '*'
            cls  = m.group(2)
            props = ' '.join(m.group(3).split()).strip('; ')
            key  = (tag, cls)
            css_map[key] = ((css_map.get(key) or '') + '; ' + props).strip('; ')

    try:
        parser  = etree.HTMLParser(recover=True)
        doc     = etree.parse(StringIO(f'<div id="_lo_wrap_">{body_str}</div>'), parser)
        root    = doc.getroot()
        wrap    = root.find('.//*[@id="_lo_wrap_"]')
        if wrap is None:
            wrap = root.find('.//body') or root

        # ── Pass 1：移除 <div title="footer"> 整塊 & <sdfield> ─────────────
        for el in list(wrap.iter()):
            if not isinstance(el.tag, str):
                continue
            tag = el.tag.lower()

            if tag == 'div' and (el.get('title') or '').lower() == 'footer':
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                continue

            if tag == 'sdfield':
                # 保留文字，移除標籤本身
                parent = el.getparent()
                if parent is not None:
                    idx = list(parent).index(el)
                    text = (el.text or '') + (el.tail or '')
                    el.tail = None
                    parent.remove(el)
                    if text:
                        if idx > 0:
                            prev = parent[idx - 1]
                            prev.tail = (prev.tail or '') + text
                        else:
                            parent.text = (parent.text or '') + text

        # ── Pass 2：<font> → <span>，保留 face/size/color ──────────────────
        _font_size_map = {
            '1': '8pt', '2': '10pt', '3': '12pt', '4': '14pt',
            '5': '18pt', '6': '24pt', '7': '36pt',
        }
        for el in list(wrap.iter('font')):
            el.tag = 'span'
            face  = el.attrib.pop('face', None)
            size  = el.attrib.pop('size', None)
            color = el.attrib.pop('color', None)

            css_parts = []
            if face:
                css_parts.append(f'font-family: {face}')
            if size:
                pt = _font_size_map.get(size.strip())
                if pt:
                    css_parts.append(f'font-size: {pt}')
            if color and color not in ('auto', '#000000', '#ffffff'):
                css_parts.append(f'color: {color}')
            if css_parts:
                existing = (el.get('style') or '').strip('; ')
                new_css  = '; '.join(css_parts)
                el.set('style', (existing + '; ' + new_css).strip('; ')
                                if existing else new_css)

        # ── Pass 3：CSS class → inline style ────────────────────────────
        for el in wrap.iter():
            if not isinstance(el.tag, str):
                continue
            tag     = el.tag.lower()
            cls_val = el.get('class', '')
            if not cls_val:
                continue
            extras = []
            for cls in cls_val.split():
                p = css_map.get((tag, cls)) or css_map.get(('*', cls))
                if p:
                    extras.append(p)
            if extras:
                existing  = (el.get('style') or '').strip('; ')
                # reversed(extras)：class 屬性中「第一個 class」排在字串末尾，
                # 使其在重複屬性時優先（CSS last-wins 語意）。
                # existing 放最後：原始 inline style 最優先（模擬 inline > class 層級）。
                new_style = '; '.join(reversed(extras))
                el.set('style', (new_style + '; ' + existing).strip('; ')
                                if existing else new_style)
            del el.attrib['class']

        # ── Pass 4：移除首個元素的 page-break-before: always ────────────
        _first_block_fixed = [False]
        for el in wrap:
            if not isinstance(el.tag, str):
                continue
            if not _first_block_fixed[0]:
                style = el.get('style', '')
                new_style = re.sub(
                    r';\s*page-break-before\s*:\s*always|page-break-before\s*:\s*always\s*;?',
                    '', style, flags=re.IGNORECASE
                ).strip('; ')
                if new_style != style:
                    el.set('style', new_style) if new_style else el.attrib.pop('style', None)
                _first_block_fixed[0] = True
                break

        # ── Pass 5：表格寬度處理 ──────────────────────────────────────────
        # 表格本身：移除絕對 px 寬度，改為 width:100%（尊重使用者邊距）
        # 欄寬：讀取 <col width="Npx">，換算成百分比後保留相對比例
        for tbl in wrap.iter('table'):
            tbl_width_px = 0
            try:
                tbl_width_px = int(tbl.get('width') or 0)
            except (ValueError, TypeError):
                pass
            tbl.attrib.pop('width', None)
            tbl.attrib.pop('cellpadding', None)
            tbl.attrib.pop('cellspacing', None)
            existing = (tbl.get('style') or '').strip('; ')
            tbl.set('style', (existing + '; width: 100%; border-collapse: collapse').strip('; ')
                             if existing else 'width: 100%; border-collapse: collapse')

            # 讀取 <col> 欄寬並換算成百分比
            cols = tbl.findall('.//col')
            if cols and tbl_width_px > 0:
                col_widths = []
                for col in cols:
                    try:
                        w = int(col.get('width') or 0)
                    except (ValueError, TypeError):
                        w = 0
                    col_widths.append(w)
                total_w = sum(col_widths)
                if total_w > 0:
                    for col, w in zip(cols, col_widths):
                        col.attrib.pop('width', None)
                        pct = round(w / total_w * 100, 1)
                        col.set('style', f'width: {pct}%')
            else:
                # 沒有 <col> 或無總寬資訊：移除個別欄的固定寬度
                for col in cols:
                    col.attrib.pop('width', None)

        # ── Pass 6：展平巢狀 <ol>（只含單個子 <ol> 的層級直接跳過）────
        def _flatten_lists(parent):
            for child in list(parent):
                if not isinstance(child.tag, str):
                    continue
                # 若 <ol> 只含一個直接子 <ol>，用內層取代外層
                if child.tag.lower() in ('ol', 'ul'):
                    real_children = [c for c in child
                                     if isinstance(c.tag, str) and c.tag.lower()
                                     not in ('', None.__class__.__name__)]
                    if (len(real_children) == 1 and
                            real_children[0].tag.lower() in ('ol', 'ul')):
                        idx = list(parent).index(child)
                        inner = real_children[0]
                        child.remove(inner)
                        parent.remove(child)
                        parent.insert(idx, inner)
                        _flatten_lists(parent)
                        return
                    _flatten_lists(child)
                else:
                    _flatten_lists(child)
        _flatten_lists(wrap)

        # ── 序列化 ────────────────────────────────────────────────────────
        out = etree.tostring(wrap, encoding='unicode', method='html')
        # 去掉外層 wrapper div
        out = re.sub(r'^<div[^>]*>', '', out)
        out = re.sub(r'</div>$', '', out)
        return out.strip()

    except Exception:
        return body_str


def _convert_ins_to_jinja(raw_bytes):
    """
    把舊系統 +++...+++ 語法全面轉換為 docxtpl Jinja2 語法。
    在段落層級操作（para.text 已拼接所有 <w:r>），安全處理 Word 的 run 切割問題。

    語法對照表：
      +++INS name+++           → {{ name }}
      +++INS $item.field+++    → {{ item.field }}
      +++FOR v IN col++++      → {% for v in col %}      （段落層級，表格外）
      +++FOR v IN $col++++     → {% tr for v in col %}   （表格列層級，表格內）
      +++END-FOR v++++         → {% endfor %}            （段落層級，表格外）
      +++END-FOR v++++         → {% tr endfor %}         （表格列層級，表格內）
    """
    import docx as python_docx

    # +++INS $item.field+++ 或 +++INS $item. field+++（允許點號前後有空格，轉換時移除）
    _INS    = re.compile(r'\+{3}INS\s+\$?\s*([A-Za-z_][\w]*(?:\s*\.\s*[A-Za-z_][\w]*)*)\s*\+{3,4}')
    # +++FOR loopVar IN $collection+++ 或 +++FOR loopVar IN collection+++
    _FOR    = re.compile(r'\+{3}FOR\s+([A-Za-z_]\w*)\s+IN\s+\$?([A-Za-z_][A-Za-z0-9_.]*)\s*\+{3,4}')
    # +++END-FOR varName+++（3 或 4 個 +）
    _ENDFOR = re.compile(r'\+{3}END-FOR\s+[A-Za-z_]\w*\s*\+{3,4}')
    # 任何 +++ 標記
    _ANY    = re.compile(r'\+{3}(?:INS|FOR|END-FOR)\s')

    doc = python_docx.Document(io.BytesIO(raw_bytes))

    def _apply(text, in_table):
        # docxtpl 識別 {%tr for %} 的 regex 是 ({%|{{)tr，無空格後才接 tr
        for_tpl    = '{%%tr for %s in %s %%}' if in_table else '{%% for %s in %s %%}'
        endfor_tpl = '{%tr endfor %}'          if in_table else '{% endfor %}'
        text = _INS.sub(lambda m: '{{ %s }}' % re.sub(r'\s', '', m.group(1)), text)
        text = _FOR.sub(lambda m: for_tpl % (m.group(1), m.group(2)), text)
        text = _ENDFOR.sub(endfor_tpl, text)
        return text

    def _process_para(para, in_table):
        if not _ANY.search(para.text):
            return
        new_text = _apply(para.text, in_table)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''

    for para in doc.paragraphs:
        _process_para(para, in_table=False)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para, in_table=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class DocEditorController(http.Controller):

    @http.route('/dobtor_doc/load', type='json', auth='user', methods=['POST'])
    def load_document(self, doc_id, **kw):
        """載入文件資料（Canvas JSON 內容＋頁面設定）。

        回傳值含 `write_date`，給前端做樂觀鎖（P2-2）：
            前端在後續 save 帶回 if_unmodified_since=write_date，
            後端比對若已變動則拒絕並回 409。
        """
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('read')
        return {
            'id': doc.id,
            'name': doc.name,
            'content_json': doc.content_json or '',
            'content_html': doc.get_content_html(),
            'header_html': doc.header_html or '',
            'footer_html': doc.footer_html or '',
            'page_format': doc.page_format,
            'margin_top': doc.margin_top,
            'margin_bottom': doc.margin_bottom,
            'margin_left': doc.margin_left,
            'margin_right': doc.margin_right,
            'model_id': doc.model_id.id if doc.model_id else False,
            'model_name': doc.model_id.model if doc.model_id else False,
            'res_id': doc.res_id or False,
            'field_aliases': doc.field_aliases or {},
            'template_field_aliases': (doc.template_id.field_aliases or {}) if doc.template_id else {},
            'template_name': doc.template_id.name if doc.template_id else '',
            'has_template': bool(doc.template_docx),
            'template_filename': doc.template_filename or '',
            'template_variables': json.loads(doc.template_variables) if doc.template_variables else [],
            'has_different_first_page': doc.has_different_first_page,
            'first_header_html': doc.first_header_html or '',
            'first_footer_html': doc.first_footer_html or '',
            # P2-2 樂觀鎖用：給前端記下最後一次同步的 write_date
            'write_date': doc.write_date.isoformat() if doc.write_date else None,
            'version_number': doc.version_number or 0,
        }

    @http.route('/dobtor_doc/save', type='json', auth='user', methods=['POST'])
    def save_document(self, doc_id=None, content_html=None, content_json=None,
                      header_html=None, footer_html=None, name=None,
                      if_unmodified_since=None, **kw):
        """儲存文件內容（content_json 為主，content_html 為備份）。

        P2-2 樂觀鎖：
            前端傳入 if_unmodified_since（從 load 回傳的 write_date），
            後端比對若 write_date 已變動 → 拒絕並回 409 形式（response 帶 conflict=True）。
            這避免兩個 user 同時編輯時，後存的人覆蓋前者沒看到的修改。

            前端收到 conflict=True 時應：
                1. 提示「文件已被他人修改」
                2. 自動 reload 拿最新內容
                3. 把使用者編輯的內容存到 IndexedDB 暫存（offline_manager）
        """
        if not doc_id:
            return {'success': False, 'error': 'doc_id required'}
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('write')

        # P2-2 樂觀鎖檢查
        if if_unmodified_since:
            current_wd = doc.write_date.isoformat() if doc.write_date else None
            # 用字串比對而非 datetime parse — 兩端都用 isoformat 應一致
            # 容忍微秒誤差：取秒為單位比對（或前端送什麼後端就比對什麼）
            if current_wd and current_wd != if_unmodified_since:
                _logger.info(
                    "Optimistic lock conflict on doc_id=%s: client had %s, server has %s",
                    doc_id, if_unmodified_since, current_wd,
                )
                return {
                    'success': False,
                    'conflict': True,
                    'error': '文件已被他人修改，請重新載入後再儲存。',
                    'server_write_date': current_wd,
                    'server_version_number': doc.version_number or 0,
                    'server_author_id': (
                        doc.write_uid.id if doc.write_uid else None
                    ),
                    'server_author_name': (
                        doc.write_uid.name if doc.write_uid else ''
                    ),
                }

        vals = {}
        if content_json is not None:
            vals['content_json'] = content_json
        if content_html is not None:
            vals['content_html'] = content_html
        if header_html is not None:
            vals['header_html'] = header_html
        if footer_html is not None:
            vals['footer_html'] = footer_html
        if name is not None:
            vals['name'] = name
        if vals:
            doc.write(vals)
        return {
            'success': True,
            'write_date': doc.write_date.isoformat(),
            'version_number': doc.version_number or 0,
        }

    # ─── DOCX 模板引擎路由 ───────────────────────────────────────────

    @http.route('/dobtor_doc/upload_template', type='http', auth='user',
                methods=['POST'], csrf=False)
    def upload_template(self, doc_id, docx_file, **kw):
        """
        上傳 DOCX 模板：
        1. 把 +++INS name+++ 語法轉換為 {{ name }}（段落層級，處理 run 切割）
        2. 偵測 {{ variable }} 佔位符清單
        3. 儲存轉換後的位元組

        Sprint 116 plus:filename 含 null byte → graceful 400(避免 Postgres
        ROLLBACK 造成 500 leak trace);filename 取 basename 防 path traversal
        傳入 DB(深度防禦原則、紀律 #15 廣域應用)。
        """
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')

        # Sprint 116 plus:filename 入口 sanitize
        raw_filename = getattr(docx_file, 'filename', '') or ''
        if '\x00' in raw_filename:
            return request.make_response(
                json.dumps({'success': False, 'error': 'invalid filename (null byte)'}),
                headers={'Content-Type': 'application/json'},
                status=400,
            )
        # 取 basename 防 path component 進 DB(深度防禦)
        safe_filename = os.path.basename(raw_filename.replace('\\', '/'))
        # 重新賦值給 docx_file.filename 讓後續邏輯用 sanitized 版本
        try:
            docx_file.filename = safe_filename
        except Exception:
            pass

        raw_bytes = docx_file.read()

        # ── Zip Bomb 防護（W1 P0-2）：檢查原檔大小、解壓總大小、entry 數量 ──
        try:
            assert_input_size(raw_bytes)
            inspect_zip_safe(raw_bytes)
        except ZipBombError as e:
            _logger.warning(
                "upload_template rejected by zip_guard: %s (file=%s, uid=%s)",
                e, getattr(docx_file, 'filename', '?'), request.env.user.id,
            )
            return request.make_response(
                json.dumps({'success': False, 'error': str(e)}),
                headers={'Content-Type': 'application/json'},
                status=400,
            )

        # 整段處理包 try/except：type='http' 路由若拋出未處理例外，Odoo 會回傳
        # HTML 500 錯誤頁，前端 resp.json() 解析失敗 → 「Unexpected token '<'」。
        # 改為一律回傳 JSON，讓真實錯誤（缺套件 / 非 docx / 解析失敗）顯示在 UI。
        variables = []
        try:
            # Step 1：轉換 +++INS+++ → {{ }}
            converted_bytes = _convert_ins_to_jinja(raw_bytes)

            # Step 2：偵測變數清單（從 ZIP 內的 document.xml）
            # 包含：{{ var }}、{{ item.field }}（取根名稱）、{% for v in collection %}（取集合名稱）
            try:
                with zipfile.ZipFile(io.BytesIO(converted_bytes)) as z:
                    xml = z.read('word/document.xml').decode('utf-8')
                # {{ var }} 或 {{ item.field }} → 取根名稱（點號前）
                raw_vars = re.findall(r'\{\{\s*([A-Za-z_][A-Za-z0-9_.]*)\s*\}\}', xml)
                root_vars = {v.split('.')[0] for v in raw_vars}
                # {% for v in collection %} 或 {% tr for v in collection %} → 取集合名稱
                for_cols = set(re.findall(
                    r'\{%-?\s+(?:tr\s+)?for\s+\w+\s+in\s+([A-Za-z_][A-Za-z0-9_.]*)\s*-?%\}', xml
                ))
                variables = sorted(root_vars | for_cols)
            except Exception:
                pass

            # Step 3：儲存轉換後的模板 + 變數清單
            doc.write({
                'template_docx': base64.b64encode(converted_bytes).decode(),
                'template_filename': docx_file.filename,
                'template_variables': json.dumps(variables),
            })
        except ImportError as e:
            _logger.exception("upload_template: 缺少 python-docx (doc_id=%s)", doc_id)
            return request.make_response(
                json.dumps({'success': False,
                            'error': '伺服器缺少 python-docx 套件，請安裝：pip install python-docx'}),
                headers={'Content-Type': 'application/json'}, status=500,
            )
        except zipfile.BadZipFile:
            return request.make_response(
                json.dumps({'success': False, 'error': '檔案不是有效的 .docx（無法解析 ZIP）'}),
                headers={'Content-Type': 'application/json'}, status=400,
            )
        except Exception as e:
            _logger.exception("upload_template 處理失敗 (doc_id=%s)", doc_id)
            return request.make_response(
                json.dumps({'success': False, 'error': '模板處理失敗：%s' % e}),
                headers={'Content-Type': 'application/json'}, status=500,
            )

        return request.make_response(
            json.dumps({'success': True, 'variables': variables}),
            headers={'Content-Type': 'application/json'}
        )

    @http.route('/dobtor_doc/fill_template', type='json', auth='user', methods=['POST'])
    def fill_template(self, doc_id, context=None, output_format='pdf', **kw):
        """
        用 docxtpl 填充模板，輸出 PDF（LibreOffice headless）或 DOCX。
        context：{variableName: value} 字典，對應模板中的 {{ variableName }}。
        """
        from docxtpl import DocxTemplate

        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')

        if not doc.template_docx:
            return {'success': False, 'error': '此文件尚未上傳 DOCX 模板'}

        raw_bytes = base64.b64decode(doc.template_docx)
        tpl = DocxTemplate(io.BytesIO(raw_bytes))
        tpl.render(context or {})

        filled_buf = io.BytesIO()
        tpl.save(filled_buf)
        filled_bytes = filled_buf.getvalue()

        if output_format == 'docx':
            return {
                'success': True,
                'content': base64.b64encode(filled_bytes).decode(),
                'filename': (doc.name or 'document') + '.docx',
                'mimetype': 'application/vnd.openxmlformats-officedocument'
                            '.wordprocessingml.document',
            }

        # PDF：LibreOffice headless
        # Sprint 70 修正：原註解誤宣稱「/usr/bin/soffice 已確認存在」，但 odoo18 container 是 minimal Ubuntu base、
        # 無 libreoffice 套件 → subprocess.run(check=True) 會 raise FileNotFoundError 500 給 user。
        # 加 graceful fallback（紀律 #11：production filesystem cross-check）。
        if not shutil.which('soffice'):
            return {
                'success': False,
                'error': 'PDF export 需要 LibreOffice — 當前 container 未安裝。請改用 output_format="docx"，'
                         '或聯絡管理員加裝 libreoffice 套件。',
                'fallback': 'docx',
            }
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, 'input.docx')
                pdf_path = os.path.join(tmpdir, 'input.pdf')
                with open(docx_path, 'wb') as f:
                    f.write(filled_bytes)
                subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'pdf',
                     '--outdir', tmpdir, docx_path],
                    check=True, timeout=60,
                    env={**os.environ, 'HOME': tmpdir}  # 防止 soffice 鎖定 ~/.config/libreoffice
                )
                with open(pdf_path, 'rb') as f:
                    pdf_bytes = f.read()
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
            return {
                'success': False,
                'error': f'LibreOffice PDF 轉換失敗：{type(e).__name__} — {str(e)[:200]}',
                'fallback': 'docx',
            }

        return {
            'success': True,
            'content': base64.b64encode(pdf_bytes).decode(),
            'filename': (doc.name or 'document') + '.pdf',
            'mimetype': 'application/pdf',
        }

    @http.route('/dobtor_doc/save_settings', type='json', auth='user', methods=['POST'])
    def save_settings(self, doc_id, **kw):
        """儲存頁面格式與邊距設定。"""
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('write')
        allowed = ('page_format', 'margin_top', 'margin_bottom',
                   'margin_left', 'margin_right',
                   'default_column_count', 'default_column_gap', 'column_rule_style')
        vals = {k: v for k, v in kw.items() if k in allowed}
        if vals:
            doc.write(vals)
        return {'success': True}

    @http.route('/dobtor_doc/fields', type='json', auth='user', methods=['POST'])
    def get_fields(self, model_name, doc_id=None, **kw):
        """取得指定模型的可用欄位清單。"""
        mixin = request.env['doc.render.mixin']
        try:
            return mixin.get_available_fields(model_name)
        except Exception as e:
            return {'error': str(e)}

    @http.route('/dobtor_doc/render_preview', type='json', auth='user', methods=['POST'])
    def render_preview(self, doc_id, record_model, record_id, **kw):
        """將欄位變數渲染為實際值（預覽用）。"""
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('read')
        try:
            record = request.env[record_model].browse(record_id)
            record.check_access_rule('read')
            rendered = doc._render_template(doc.get_content_html(), record, with_chip=True)
            return {'html': rendered}
        except Exception as e:
            return {'error': str(e)}

    # ─── 中文 token 別名（L2-v2 alias map）─────────────────────────────
    @http.route('/dobtor_doc/aliases/get', type='json', auth='user', methods=['POST'])
    def get_aliases(self, doc_id, **kw):
        """讀取文件目前的中文 token → Jinja2 expression 對映。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')
        return {'aliases': doc.field_aliases or {}}

    @http.route('/dobtor_doc/aliases/save', type='json', auth='user', methods=['POST'])
    def save_aliases(self, doc_id, aliases, **kw):
        """整批覆寫文件的中文 token → Jinja2 expression 對映。

        aliases: dict[str, str]，key=中文 token、value=Jinja2 expression（不含 {{ }}）。
        例：{"工程名稱": "object.project_id.name"}
        """
        if not isinstance(aliases, dict):
            return {'error': 'aliases 必須為 dict'}
        # 防呆：剝除空 key / 空 value、key 移除前後 《》
        cleaned = {}
        for raw_key, raw_val in aliases.items():
            if not raw_key or not raw_val:
                continue
            key = str(raw_key).strip().strip('《》').strip()
            val = str(raw_val).strip()
            if key and val:
                cleaned[key] = val
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        doc.write({'field_aliases': cleaned})
        return {'success': True, 'aliases': cleaned}

    @http.route('/dobtor_doc/aliases/auto_init', type='json', auth='user', methods=['POST'])
    def auto_init_aliases(self, doc_id, overwrite=False, **kw):
        """從 doc.model_id 的欄位自動生成中文 alias 對映。

        overwrite=False（預設）保留既有 token；True 整批以模型欄位重建。
        """
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        return doc.init_aliases_from_model(overwrite=bool(overwrite))

    @http.route('/dobtor_doc/aliases/scan_convert', type='json', auth='user', methods=['POST'])
    def scan_convert_aliases(self, doc_id, **kw):
        """掃描文件內所有 {{ expression }} 文字，根據既有 alias map 反查中文 token 後改寫成 《token》。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        return doc.scan_and_convert_to_alias()

    @http.route('/dobtor_doc/template_aliases/get', type='json', auth='user', methods=['POST'])
    def get_template_aliases(self, doc_id, **kw):
        """讀取 doc 所屬 template 的 alias map（範本層級全域對映）。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')
        if not doc.template_id:
            return {'aliases': {}, 'template_id': False, 'template_name': ''}
        return {
            'aliases': doc.template_id.field_aliases or {},
            'template_id': doc.template_id.id,
            'template_name': doc.template_id.name or '',
        }

    @http.route('/dobtor_doc/template_aliases/save', type='json', auth='user', methods=['POST'])
    def save_template_aliases(self, doc_id, aliases, **kw):
        """覆寫 doc 所屬 template 的 alias map。所有使用此範本的 doc 都會生效。"""
        if not isinstance(aliases, dict):
            return {'error': 'aliases 必須為 dict'}
        cleaned = {}
        for raw_k, raw_v in aliases.items():
            if not raw_k or not raw_v:
                continue
            k = str(raw_k).strip().strip('《》').strip()
            v = str(raw_v).strip()
            if k and v:
                cleaned[k] = v
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        if not doc.template_id:
            return {'error': '此文件未綁定範本，無法寫入範本級 alias'}
        doc.template_id.check_access_rule('write')
        doc.template_id.write({'field_aliases': cleaned})
        return {'success': True, 'aliases': cleaned}

    @http.route('/dobtor_doc/preview_content_json', type='json', auth='user', methods=['POST'])
    def preview_content_json(self, doc_id, record_id=None, content_json=None, **kw):
        """把 content_json 內所有 token / {{ var }} 替換成實際值，回傳新的 JSON。

        編輯器「預覽模式」用：toggle 切到預覽時暫存原 content_json，呼叫本端點取代渲染後的版本，
        canvas-editor 直接 executeSetValue 上去，使用者可以看見實際值。

        content_json：可選。給了就渲染「這份當前編輯器內容」（保留已建的 control chip，只把
        殘餘 token 文字換成值）；沒給則用 doc 儲存的 content_json。前端開檔升級 chip 後會把
        當前內容傳進來，讓 chip 與其餘 token 的實際值「共存」而非被整份覆蓋。
        """
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')
        rec_id = int(record_id) if record_id else doc.res_id
        if not doc.model_id or not rec_id:
            return {'error': '此文件未綁定 model_id 或 res_id'}
        try:
            record = request.env[doc.model_id.model].browse(rec_id)
            record.check_access_rule('read')
            if not record.exists():
                return {'error': f'記錄 {doc.model_id.model}/{rec_id} 不存在'}
        except Exception as e:
            return {'error': f'記錄存取失敗：{e}'}

        # 合併 alias map（template + doc 自己）
        aliases = doc._collect_field_aliases()
        if not aliases:
            return {'error': '尚無 alias 對映'}

        # 拆 token / varname
        import re as _re
        token_alias = {}
        var_alias = {}
        for k, v in aliases.items():
            ks = str(k).strip()
            vs = str(v).strip()
            if not ks or not vs:
                continue
            if _re.match(r'^[A-Za-z_]\w*$', ks):
                var_alias[ks] = vs
            else:
                token_alias[ks] = vs

        # 準備 Jinja2 環境（含 helper）
        from jinja2.sandbox import SandboxedEnvironment
        env_j = SandboxedEnvironment()
        for name, fn in doc._get_render_helpers(record).items():
            env_j.globals[name] = fn

        def _eval(expression):
            try:
                tpl = env_j.from_string('{{ ' + expression + ' }}')
                return tpl.render(object=record, user=request.env.user)
            except Exception:
                return ''

        # 預編譯所有 alias 的渲染結果（cache 重複 token）
        token_to_value = {tk: _eval(expr) for tk, expr in token_alias.items()}
        var_to_value = {vn: _eval(expr) for vn, expr in var_alias.items()}

        _token_pat = _re.compile(r'《([^》]+)》')
        _var_pat = _re.compile(r'\{\{\s*([A-Za-z_]\w*)\s*\}\}')

        def _replace_text(text):
            if not isinstance(text, str) or not text:
                return text
            def _rt(m):
                return token_to_value.get(m.group(1).strip(), m.group(0))
            text = _token_pat.sub(_rt, text)
            def _rv(m):
                return var_to_value.get(m.group(1).strip(), m.group(0))
            return _var_pat.sub(_rv, text)

        # 遞迴掃 content_json：優先用前端傳入的當前內容（含 chip），否則用 doc 儲存的
        import json as _json
        cj = content_json if content_json is not None else doc.content_json
        if isinstance(cj, str):
            try:
                cj = _json.loads(cj)
            except Exception:
                return {'error': 'content_json 無法解析'}
        if not cj:
            return {'error': 'content_json 為空'}

        # 深複製避免改到 cache
        import copy as _copy
        cj = _copy.deepcopy(cj)

        def _walk(node):
            if isinstance(node, dict):
                for k, v in node.items():
                    if k == 'value' and isinstance(v, str):
                        node[k] = _replace_text(v)
                    else:
                        _walk(v)
            elif isinstance(node, list):
                for it in node:
                    _walk(it)

        _walk(cj)
        return {'success': True, 'content_json': cj}

    @http.route('/dobtor_doc/preview/<int:doc_id>', type='http', auth='user', methods=['GET'])
    def preview_document(self, doc_id, record_id=None, **kw):
        """瀏覽器直接打開的預覽頁（form view「快速預覽」按鈕用）。

        參數：
          record_id：可選；未指定時用 doc.res_id。
        """
        doc = request.env['doc.document'].browse(int(doc_id))
        try:
            doc.check_access_rule('read')
        except Exception:
            return request.not_found()

        # 決定渲染用 record
        rec_id = int(record_id) if record_id else doc.res_id
        rec = None
        rec_model = doc.model_id.model if doc.model_id else None
        if rec_model and rec_id:
            try:
                rec = request.env[rec_model].browse(rec_id)
                rec.check_access_rule('read')
                if not rec.exists():
                    rec = None
            except Exception:
                rec = None

        # 渲染（帶 chip 樣式）
        body = doc._render_template(doc.get_content_html(), rec, with_chip=True) if rec \
            else doc.get_content_html() or '<p><em>（文件無內容）</em></p>'

        doc_name = html_mod.escape(doc.name or '未命名文件')
        record_label = html_mod.escape(rec.display_name) if rec else '（無綁定記錄）'
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>預覽：{doc_name}</title>
<style>
body {{
  font-family: 'Microsoft JhengHei', 'Noto Sans TC', Arial, sans-serif;
  padding: 24px; max-width: 820px; margin: auto;
  background: #f8fafc;
}}
.doc-preview-header {{
  background: #fff;
  border: 1px solid #e2e8f0;
  border-radius: 6px;
  padding: 10px 14px;
  margin-bottom: 16px;
  font-size: 13px;
  color: #475569;
  display: flex;
  justify-content: space-between;
  flex-wrap: wrap;
  gap: 8px;
}}
.doc-preview-body {{
  background: #fff;
  padding: 24px;
  border: 1px solid #e2e8f0;
  border-radius: 6px;
}}
@media print {{
  body {{ padding: 0; max-width: none; background: #fff; }}
  .doc-preview-header {{ display: none; }}
  .doc-preview-body {{ border: 0; padding: 0; }}
  .doc-field-token {{ background: transparent; border: 0; padding: 0; color: inherit; }}
}}
.doc-field-token {{
  background: #e3f2fd;
  border: 1px solid #90caf9;
  border-radius: 3px;
  padding: 1px 4px;
  color: #1565c0;
  font-size: 0.95em;
}}
.doc-field-token:empty::after {{ content: '（無值）'; color: #999; font-style: italic; }}
</style></head><body>
<div class="doc-preview-header">
  <span><strong>{doc_name}</strong> 預覽</span>
  <span>對應記錄：{record_label}</span>
</div>
<div class="doc-preview-body">{body}</div>
</body></html>"""
        return request.make_response(
            full_html,
            headers=[('Content-Type', 'text/html; charset=utf-8')],
        )

    @http.route('/dobtor_doc/template_aliases/auto_init', type='json', auth='user', methods=['POST'])
    def auto_init_template_aliases(self, doc_id, overwrite=False, **kw):
        """從 doc 綁定 model 自動生成範本級 alias，寫入 doc.template_id。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        if not doc.template_id:
            return {'success': False, 'error': '此文件未綁定範本'}
        if not doc.model_id:
            return {'success': False, 'error': '此文件未綁定 Odoo 模型'}
        # 借用 doc 的 init_aliases_from_model 邏輯，但寫入點改為 template
        tmpl = doc.template_id
        tmpl.check_access_rule('write')
        # 暫借 doc 的方法計算，再把結果搬到 template
        # 用 _new() 避免污染現有 doc.field_aliases；直接呼叫靜態生成
        return tmpl.init_aliases_from_model_for(
            doc.model_id.model,
            overwrite=bool(overwrite),
        )

    @http.route('/dobtor_doc/export', type='json', auth='user', methods=['POST'])
    def export_document(self, doc_id, format='pdf', quality='high',
                        record_model=None, record_id=None, **kw):
        """匯出文件為 PDF 或 DOCX。quality: 'high'（後端）"""
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('read')

        record = None
        if record_model and record_id:
            try:
                record = request.env[record_model].browse(record_id)
                record.check_access_rule('read')
            except Exception:
                record = None

        # Fallback：若前端沒帶 record，自動從 doc.model_id + doc.res_id 取
        # （讓「即用即匯出」流程不必每次都記得帶綁定資訊）
        if record is None and doc.model_id and doc.res_id:
            try:
                rec = request.env[doc.model_id.model].browse(doc.res_id)
                rec.check_access_rule('read')
                if rec.exists():
                    record = rec
            except Exception:
                record = None

        try:
            if format == 'pdf':
                file_bytes = doc._generate_pdf(record=record)
                mimetype = 'application/pdf'
                filename = f'{doc.name}.pdf'
            elif format == 'docx':
                file_bytes = doc._generate_docx_via_libreoffice(record=record)
                mimetype = ('application/vnd.openxmlformats-officedocument'
                            '.wordprocessingml.document')
                filename = f'{doc.name}.docx'
            else:
                return {'error': f'不支援的格式：{format}'}

            request.env['doc.editor.export.log'].record_export(
                doc, format, with_alias=record is not None,
                file_size=len(file_bytes), source='editor',
            )
            return {
                'filename': filename,
                'data': base64.b64encode(file_bytes).decode(),
                'mimetype': mimetype,
            }
        except Exception as e:
            return {'error': str(e)}

    @http.route('/dobtor_doc/save_version', type='json', auth='user', methods=['POST'])
    def save_version(self, doc_id, label=None, **kw):
        """儲存版本快照（W7-8 P1-1：回傳 version_number 與 message_id）。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        result = doc.action_save_version(label=label)
        return {'success': True, **(result or {})}

    # ─── 版本管理路由（W7-8 P1-1）────────────────────────────────────

    @http.route('/dobtor_doc/versions/list', type='json', auth='user', methods=['POST'])
    def versions_list(self, doc_id, **kw):
        """列出文件所有版本快照（不含 content，輕量）。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')
        return {'versions': doc.get_version_list()}

    @http.route('/dobtor_doc/versions/get', type='json', auth='user', methods=['POST'])
    def versions_get(self, doc_id, version_id=None, message_id=None, **kw):
        """取得單一版本的完整內容（accept version_id 或 legacy message_id）。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')
        vid = version_id if version_id is not None else message_id
        content = doc.get_version_content(vid)
        if content is None:
            return {'error': '找不到指定的版本快照'}
        return {'success': True, **content}

    @http.route('/dobtor_doc/versions/restore', type='json', auth='user', methods=['POST'])
    def versions_restore(self, doc_id, version_id=None, message_id=None, **kw):
        """還原到指定版本（自動先存「還原前」快照）。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('write')
        vid = version_id if version_id is not None else message_id
        try:
            result = doc.restore_version(vid)
        except Exception as e:
            return {'error': str(e)}
        return {'success': True, **(result or {})}

    @http.route('/dobtor_doc/versions/diff', type='json', auth='user', methods=['POST'])
    def versions_diff(self, doc_id, version_id_a=None, version_id_b=None,
                      message_id_a=None, message_id_b=None, **kw):
        """段落層級 diff 兩個版本（accept version_id_* 或 legacy message_id_*）。"""
        doc = request.env['doc.document'].browse(int(doc_id))
        doc.check_access_rule('read')
        a = version_id_a if version_id_a is not None else message_id_a
        b = version_id_b if version_id_b is not None else message_id_b
        result = doc.diff_versions(a, b)
        if result is None:
            return {'error': '找不到指定的版本快照'}
        return {'success': True, **result}

    # ─── 監控與遙測（P2-4）─────────────────────────────────────────

    @http.route('/dobtor_doc/telemetry/error', type='json', auth='user', methods=['POST'])
    def telemetry_error(self, error_type='other', message='', stack_trace='',
                        user_agent='', url='', doc_id=None, extra=None, **kw):
        """前端錯誤上報。

        受信賴的最小欄位 — 拒絕儲存超過 size 的 message / stack 以防 abuse。
        """
        ALLOWED_TYPES = (
            'js_error', 'promise_rejection', 'canvas_error',
            'save_failure', 'import_failure', 'export_failure', 'other',
        )
        if error_type not in ALLOWED_TYPES:
            error_type = 'other'

        # 截斷防止濫用（DB 欄位本身 size=512/256，但前端可能傳更長）
        message = (message or '')[:500]
        user_agent = (user_agent or '')[:250]
        url = (url or '')[:500]
        # stack_trace 是 Text 沒長度限制，但截到 8K 避免 DB 爆量
        stack_trace = (stack_trace or '')[:8000]

        try:
            request.env['doc.editor.error.log'].sudo().create({
                'doc_id': int(doc_id) if doc_id else False,
                'user_id': request.env.user.id,
                'company_id': request.env.company.id,
                'error_type': error_type,
                'message': message,
                'stack_trace': stack_trace,
                'user_agent': user_agent,
                'url': url,
                'extra': extra or {},
            })
            return {'success': True}
        except Exception as e:
            _logger.warning("Failed to log telemetry error: %s", e)
            # 不擾斷前端，靜默吞掉（telemetry 失敗不應影響使用者）
            return {'success': False}

    @http.route('/dobtor_doc/telemetry/metric', type='json', auth='user', methods=['POST'])
    def telemetry_metric(self, metric_type=None, value=None,
                         doc_id=None, page_count=None, extra=None, **kw):
        """前端效能指標上報。

        每筆代表一個 metric event。前端可批次合併多個 metric type 一次發。
        """
        if not metric_type or value is None:
            return {'success': False, 'error': 'metric_type + value required'}
        # metric_type 不限白名單（彈性高），但截 size 避免 abuse
        metric_type = str(metric_type)[:64]

        try:
            value_f = float(value)
        except (TypeError, ValueError):
            return {'success': False, 'error': 'value must be numeric'}

        try:
            request.env['doc.editor.perf.metric'].sudo().create({
                'doc_id': int(doc_id) if doc_id else False,
                'user_id': request.env.user.id,
                'metric_type': metric_type,
                'value': value_f,
                'page_count': int(page_count) if page_count else 0,
                'extra': extra or {},
            })
            return {'success': True}
        except Exception as e:
            _logger.warning("Failed to log telemetry metric: %s", e)
            return {'success': False}

    @http.route('/dobtor_doc/import', type='http', auth='user', methods=['POST'], csrf=False)
    def import_document(self, **kw):
        """匯入 DOCX / ODT 檔案。

        Engine 並行通道（Phase E）：
            engine=libreoffice（預設）：傳統 LibreOffice → HTML 路徑（穩定）
            engine=ts                 ：本模組 TS OOXML Parser → IElement[] 路徑
                                        前端可直接餵給 canvas-editor 初始化（content_json）
            engine=both               ：兩條都跑，回傳 html + elements + 比對 log
                                        debug 模式：可快速肉眼比對 LibreOffice 與 TS 路徑差異

        前端策略（doc_editor.js）：
            - 優先使用 elements（如有），呼叫 editor.command.executeSetValue(elements)
            - 否則 fallback 到 html（既有路徑）

        ODT / 其他格式仍只能走 LibreOffice。

        Args:
            engine: 'libreoffice' / 'ts' / 'both'（從 form 或 query string 取）
        """
        def _json_resp(data):
            return request.make_response(
                json.dumps(data, ensure_ascii=False),
                headers={'Content-Type': 'application/json; charset=utf-8'},
            )

        upload = request.httprequest.files.get('file')
        if not upload:
            return _json_resp({'error': '未收到檔案'})

        filename = upload.filename or ''
        ext = os.path.splitext(filename)[1].lower()

        # 解析 engine 參數（form > query > 預設）
        engine = (
            request.httprequest.form.get('engine')
            or request.httprequest.args.get('engine')
            or 'libreoffice'
        ).lower()
        if engine not in ('libreoffice', 'ts', 'both'):
            engine = 'libreoffice'

        # ODT 不支援 TS 路徑（無 ODT parser），自動降級
        if ext == '.odt' and engine in ('ts', 'both'):
            engine = 'libreoffice'

        # Sprint Y58：opt-in flag（form > query > 預設 false）。
        # 預設值維持與 Sprint 358-359 後的行為一致 — 不啟用 floatTextBox 展平、
        # 不透傳 wp:anchor 屬性。caller 想要時送 `float_textbox=1` / `anchored_image=1`。
        def _truthy(val):
            return str(val or '').strip().lower() in ('1', 'true', 'yes', 'on')

        float_textbox = _truthy(
            request.httprequest.form.get('float_textbox')
            or request.httprequest.args.get('float_textbox')
        )
        anchored_image = _truthy(
            request.httprequest.form.get('anchored_image')
            or request.httprequest.args.get('anchored_image')
        )

        try:
            file_bytes = upload.read()

            # ── Zip Bomb 防護（W1 P0-2）：DOCX 才檢查；ODT 也是 zip 但結構不同 ──
            if ext in ('.docx', '.odt'):
                try:
                    assert_input_size(file_bytes)
                    inspect_zip_safe(file_bytes)
                except ZipBombError as e:
                    _logger.warning(
                        "import_document rejected by zip_guard: %s (file=%s, uid=%s)",
                        e, filename, request.env.user.id,
                    )
                    return _json_resp({'error': str(e)})

            page_margins = None
            body_html = None
            elements = None
            audit = {}  # debug 比對資訊

            # ── TS 路徑（engine=ts 或 both）──
            if engine in ('ts', 'both') and ext == '.docx':
                ts_elements = _ts_parse_docx_to_elements(
                    file_bytes,
                    float_textbox=float_textbox,
                    anchored_image=anchored_image,
                )
                if ts_elements is not None:
                    elements = ts_elements
                    audit['ts_element_count'] = len(ts_elements)
                else:
                    audit['ts_failed'] = True
                    if engine == 'ts':
                        # 純 ts 模式失敗時自動 fallback libreoffice（避免使用者卡住）
                        engine = 'libreoffice'

            # ── LibreOffice 路徑（engine=libreoffice 或 both）──
            if engine in ('libreoffice', 'both'):
                if ext in ('.docx', '.odt'):
                    lo_result = _lo_convert_to_html(file_bytes, ext)
                    if lo_result is not None:
                        body_html, page_margins = lo_result
                        audit['lo_html_len'] = len(body_html)
                    elif ext == '.docx':
                        body_html = _docx_to_html_with_format(file_bytes)
                        audit['lo_fallback'] = 'docx_python'
                    else:
                        body_html = _odt_to_html(file_bytes)
                        audit['lo_fallback'] = 'odt_python'
                else:
                    if not shutil.which('soffice'):
                        return _json_resp({
                            'error': f'不支援的格式（{ext}）。支援格式：.docx、.odt'
                        })
                    lo_result = _lo_convert_to_html(file_bytes, ext)
                    if lo_result is None:
                        return _json_resp({'error': f'LibreOffice 無法轉換格式：{ext}'})
                    body_html, page_margins = lo_result

            resp = {'engine': engine}
            if body_html is not None:
                resp['html'] = body_html
            if elements is not None:
                resp['elements'] = elements
            if page_margins:
                resp['margins'] = page_margins
            if engine == 'both':
                resp['audit'] = audit

            # 至少要有一條路徑成功
            if body_html is None and elements is None:
                return _json_resp({
                    'error': 'TS 與 LibreOffice 皆無法轉換此檔案',
                    'audit': audit,
                })

            return _json_resp(resp)

        except Exception as e:
            return _json_resp({'error': str(e)})

    # ─── Phase F：視覺回歸測試路由 ─────────────────────────────────────────
    #
    # 用途：給 puppeteer + pixelmatch 截圖比對 LibreOffice golden PNG 用。
    # 設計：
    #   GET  /dobtor_doc_editor/test?fixture=<rel_path>  → 回傳 clean HTML 頁面
    #   POST /dobtor_doc_editor/test_data {fixture}       → 回傳 IElement[]
    #   test_harness.js 在 clean 頁面 fetch test_data → 餵 canvas-editor → 設 ready flag

    @http.route('/dobtor_doc_editor/test', type='http', auth='user', methods=['GET'])
    def test_render(self, fixture=None, **kw):
        """渲染 fixture .docx 為 clean canvas-editor 頁面（無 Odoo header/sidebar）。

        Phase F 視覺回歸 pipeline 入口。puppeteer 對此 URL 截圖，
        對比 tests/fixtures/<category>/golden/<fixture>-<page>.png。
        """
        if not fixture:
            return request.make_response('missing ?fixture=<rel_path>', status=400)

        # 安全：fixture 必須是 tests/fixtures/ 下的相對路徑、無 .. traversal
        module_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        fixtures_root = os.path.join(module_dir, 'tests', 'fixtures')
        abs_path = os.path.normpath(os.path.join(fixtures_root, fixture))
        if not abs_path.startswith(fixtures_root + os.sep):
            return request.make_response('invalid fixture path', status=400)
        if not abs_path.lower().endswith('.docx'):
            return request.make_response('only .docx supported', status=400)
        if not os.path.isfile(abs_path):
            return request.make_response(f'fixture not found: {fixture}', status=404)

        return request.render('dobtor_doc_editor.test_layout', {
            'fixture_name': fixture,
        })

    @http.route('/dobtor_doc_editor/test_data', type='json', auth='user', methods=['POST'])
    def test_data(self, fixture=None, float_textbox=False, anchored_image=False, **kw):
        """回傳指定 fixture 的 IElement[]（Phase F test_harness.js 用）。

        參數：
            fixture:        tests/fixtures/ 下的相對路徑（如 '01_simple/xxx.docx'）
            float_textbox:  Sprint Y58 opt-in 展平 wp:anchor + w:txbxContent 文字
            anchored_image: Sprint Y58 opt-in 透傳 wp:anchor 屬性

        回傳：{'elements': [...IElement...]} 或 {'error': str}
        """
        if not fixture:
            return {'error': 'missing fixture parameter'}

        module_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        fixtures_root = os.path.join(module_dir, 'tests', 'fixtures')
        abs_path = os.path.normpath(os.path.join(fixtures_root, fixture))
        if not abs_path.startswith(fixtures_root + os.sep):
            return {'error': 'invalid fixture path'}
        if not abs_path.lower().endswith('.docx') or not os.path.isfile(abs_path):
            return {'error': f'fixture not found: {fixture}'}

        with open(abs_path, 'rb') as fp:
            file_bytes = fp.read()

        elements = _ts_parse_docx_to_elements(
            file_bytes,
            float_textbox=bool(float_textbox),
            anchored_image=bool(anchored_image),
        )
        if elements is None:
            return {'error': 'TS parser failed (CLI not built or runtime error)'}

        return {'elements': elements, 'fixture': fixture}

    # ════════════════════════════════════════════════════════════════
    # Phase 8 Template UI Builder（ADR-022）— Phase 2.1 inline control
    # ════════════════════════════════════════════════════════════════
    #
    # 設計：doc-centric endpoint。前端傳 doc_id、後端反查 doc.template_id 操作。
    # 為什麼：1) 跟既有 /dobtor_doc/* 端點命名一致；2) 用 doc 的 access rule
    #         自然繼承（user 對 doc 有 write 才能改範本欄位）；3) 不需要前端管 template_id。
    #
    # 權限：load read-only / save/delete 需要 doc.write（manager group 才能寫 doc.template.field）。
    # 不 sudo() — 範本欄位修改 = 範本設計變更，依 ACL 限定 manager 是有意為之。

    @http.route('/dobtor_doc/template_fields/load', type='json', auth='user', methods=['POST'])
    def template_fields_load(self, doc_id, **kw):
        """載入當前 doc 對應 template 的所有 signers + fields。

        回傳：
            {
                'has_template': bool,
                'template_id': int | None,
                'signers': [{id, name, color, sequence, field_count}, ...],
                'fields':  [{id, signer_id, field_type, page_no, required,
                             placeholder_text, font_size, odoo_field_name,
                             width, height, pos_x, pos_y}, ...]
            }
        """
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('read')
        if not doc.template_id:
            return {'has_template': False, 'template_id': None, 'signers': [], 'fields': []}
        template = doc.template_id
        template.check_access_rule('read')
        signers = template.signer_ids.read([
            'id', 'name', 'color', 'sequence', 'field_count',
        ])
        fields_list = template.field_ids.read([
            'id', 'signer_id', 'field_type', 'page_no', 'required',
            'placeholder_text', 'font_size', 'odoo_field_name',
            'width', 'height', 'pos_x', 'pos_y',
            'layout_mode',  # Sprint D
            # 互動式 control 選項設定
            'option_source', 'selection_field_name', 'input_able', 'is_multi_select',
        ])
        # signer_id 從 Odoo Many2one [id, display_name] tuple 簡化為純 id
        # 並附上自訂選項清單（option_ids → options）供前端組 valueSets
        field_recs = {rec.id: rec for rec in template.field_ids}
        for f in fields_list:
            if f.get('signer_id'):
                f['signer_id'] = f['signer_id'][0]
            rec = field_recs.get(f['id'])
            f['options'] = [
                {'value': o.value, 'code': o.code or o.value}
                for o in rec.option_ids.sorted('sequence')
            ] if rec else []
        return {
            'has_template': True,
            'template_id': template.id,
            'signers': signers,
            'fields': fields_list,
        }

    def _field_control_spec(self, field, record, aliases=None):
        """把一個 doc.template.field 轉成前端 canvas-editor control 所需的設定。

        record：doc 綁定的 Odoo record（可能為 None）；用來算 current_code
                （chip 開啟時的預設選中值＝該 record 該欄位的當前值）。
        aliases：doc._collect_field_aliases()，用來算此欄位對應的文件內 token 字面字串
                （含 {{ varname }} 與 《中文》兩種格式），供前端在內容中尋找並升級。

        odoo 來源優先用 record 的 model 解析 Selection（拿得到 current_code）；
        沒 record 時退而用 template.model_id，只給選項、不給預設值。
        """
        value_sets = []
        current_code = None
        if field.option_source == 'odoo' and field.selection_field_name:
            model_name = None
            if record is not None and record.exists():
                model_name = record._name
            elif field.template_id.model_id:
                model_name = field.template_id.model_id.model
            if model_name and model_name in request.env:
                fdef = request.env[model_name]._fields.get(field.selection_field_name)
                if fdef and fdef.type == 'selection':
                    for tech, label in fdef._description_selection(request.env):
                        value_sets.append({'value': label, 'code': tech})
                    if (record is not None and record.exists()
                            and field.selection_field_name in record._fields):
                        val = record[field.selection_field_name]
                        if val:
                            current_code = val
        elif field.option_source == 'custom':
            for opt in field.option_ids.sorted('sequence'):
                value_sets.append({'value': opt.value, 'code': opt.code or opt.value})

        # 此欄位對應的文件內 token 字面字串：以 placeholder_text（varname）為錨，
        # 找 field_aliases 中對映到同一 expression 的所有 key。
        # varname-like key → {{ key }}；中文等非 varname key → 《key》。
        # 讓前端能同時升級 {{ timing }}（新文件）與 《檢查時機》（已遷移文件）。
        varname = field.placeholder_text or ''
        tokens = []
        if varname:
            tokens.append('{{ %s }}' % varname)
            if aliases:
                expr = aliases.get(varname)
                if expr:
                    for k, v in aliases.items():
                        if v != expr or k == varname:
                            continue
                        if re.match(r'^[A-Za-z_]\w*$', k):
                            tokens.append('{{ %s }}' % k)
                        else:
                            tokens.append('《%s》' % k)
        tokens = list(dict.fromkeys(tokens))  # 去重、保序

        return {
            'field_id': field.id,
            'placeholder_text': varname,
            'control_type': field.field_type,
            'value_sets': value_sets,
            'input_able': field.input_able,
            'is_multi_select': field.is_multi_select,
            'current_code': current_code,
            'concept_id': str(field.id),
            'tokens': tokens,
        }

    @http.route('/dobtor_doc/template_fields/options', type='json', auth='user', methods=['POST'])
    def template_field_options(self, doc_id, field_id=None, **kw):
        """回傳互動式 control 的選項設定（valueSets + 預設值）。

        給 field_id → 回單一欄位 spec；不給 → 回此文件範本所有欄位的 spec（批次，
        供「開文件自動升級」一次拿齊，免逐個 round-trip）。
        """
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('read')
        record = doc._resolve_bound_record()
        aliases = doc._collect_field_aliases() or {}
        FieldModel = request.env['doc.template.field']
        if field_id:
            field = FieldModel.browse(int(field_id))
            if not field.exists() or (doc.template_id and field.template_id != doc.template_id):
                return {'success': False, 'error': '欄位不存在或不屬於此範本'}
            return {'success': True, 'spec': self._field_control_spec(field, record, aliases)}
        specs = []
        if doc.template_id:
            for field in doc.template_id.field_ids:
                specs.append(self._field_control_spec(field, record, aliases))
        return {'success': True, 'specs': specs}

    @http.route('/dobtor_doc/template_fields/save_field', type='json', auth='user', methods=['POST'])
    def template_fields_save_field(self, doc_id, field=None, **kw):
        """新建或更新單個範本欄位。

        field 參數：
            {
                'id':            int | None  # None = create
                'signer_id':     int (required)
                'field_type':    str
                'page_no':       int
                'required':      bool
                'placeholder_text': str
                'font_size':     int
                'odoo_field_name': str
                'width', 'height', 'pos_x', 'pos_y': float
            }

        回傳：{'success': True, 'id': field_id, 'signer_field_counts': {signer_id: count}}
            或 {'success': False, 'error': str}
        """
        if not field:
            return {'success': False, 'error': 'missing field parameter'}
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('write')
        if not doc.template_id:
            return {'success': False, 'error': '此文件未關聯範本，無法新增欄位'}
        template = doc.template_id
        # 允許的欄位白名單（防止前端塞奇怪 key 進來）
        ALLOWED = {
            'signer_id', 'field_type', 'page_no', 'required',
            'placeholder_text', 'font_size', 'odoo_field_name',
            'width', 'height', 'pos_x', 'pos_y',
            'layout_mode',  # Sprint D
            # 互動式 control 選項設定
            'option_source', 'selection_field_name', 'input_able', 'is_multi_select',
        }
        vals = {k: v for k, v in field.items() if k in ALLOWED}
        FieldModel = request.env['doc.template.field']
        field_id = field.get('id')
        try:
            if field_id:
                rec = FieldModel.browse(int(field_id))
                rec.check_access_rule('write')
                # 確保不能透過 update 把欄位搬到其他範本
                if rec.template_id != template:
                    return {'success': False, 'error': '欄位不屬於此範本'}
                rec.write(vals)
            else:
                vals['template_id'] = template.id
                rec = FieldModel.create(vals)
                field_id = rec.id
        except Exception as e:
            return {'success': False, 'error': str(e)}

        # 回傳每位簽約人最新的欄位計數（前端 chip 計數即時更新）
        signer_counts = {}
        for signer in template.signer_ids:
            signer_counts[signer.id] = len(signer.field_ids)

        return {
            'success': True,
            'id': field_id,
            'signer_field_counts': signer_counts,
            'field_count': len(template.field_ids),
        }

    @http.route('/dobtor_doc/template_fields/delete_field', type='json', auth='user', methods=['POST'])
    def template_fields_delete_field(self, doc_id, field_id, **kw):
        """刪除單個範本欄位。

        回傳：{'success': True, 'signer_field_counts': {...}, 'field_count': int}
            或 {'success': False, 'error': str}
        """
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('write')
        if not doc.template_id:
            return {'success': False, 'error': '此文件未關聯範本'}
        template = doc.template_id
        field = request.env['doc.template.field'].browse(int(field_id))
        if not field.exists():
            return {'success': False, 'error': 'field not found'}
        if field.template_id != template:
            return {'success': False, 'error': '欄位不屬於此範本'}
        try:
            field.check_access_rule('unlink')
            field.unlink()
        except Exception as e:
            return {'success': False, 'error': str(e)}

        signer_counts = {}
        for signer in template.signer_ids:
            signer_counts[signer.id] = len(signer.field_ids)
        return {
            'success': True,
            'signer_field_counts': signer_counts,
            'field_count': len(template.field_ids),
        }

    @http.route('/dobtor_doc/template_fields/save_signer', type='json', auth='user', methods=['POST'])
    def template_fields_save_signer(self, doc_id, signer=None, **kw):
        """新建或更新範本簽約人（讓 Phase 2.1 UI 也能在文件編輯器加 signer）。

        signer 參數：
            {'id': int | None, 'name': str, 'color': int, 'sequence': int}
        """
        if not signer:
            return {'success': False, 'error': 'missing signer parameter'}
        doc = request.env['doc.document'].browse(doc_id)
        doc.check_access_rule('write')
        if not doc.template_id:
            return {'success': False, 'error': '此文件未關聯範本'}
        template = doc.template_id
        ALLOWED = {'name', 'color', 'sequence'}
        vals = {k: v for k, v in signer.items() if k in ALLOWED}
        SignerModel = request.env['doc.template.signer']
        signer_id = signer.get('id')
        try:
            if signer_id:
                rec = SignerModel.browse(int(signer_id))
                rec.check_access_rule('write')
                if rec.template_id != template:
                    return {'success': False, 'error': '簽約人不屬於此範本'}
                rec.write(vals)
            else:
                vals['template_id'] = template.id
                rec = SignerModel.create(vals)
                signer_id = rec.id
        except Exception as e:
            return {'success': False, 'error': str(e)}
        return {'success': True, 'id': signer_id}

    # ──────────────────────────────────────────────────────────────────────
    # Sprint A —— Sub-nav 預覽 / 請求清單端點
    #
    # 設計：兩個端點都 doc-centric（前端傳 doc_id）；預覽端點用 jinja2 sandbox
    # 渲染 doc.content_html，避開既有 fill_template 對 LibreOffice 的硬依賴。
    # 請求端點目前回傳殼資料；未來接上 doc.fill.request model 後改回真實清單。
    # ──────────────────────────────────────────────────────────────────────

    @http.route('/dobtor_doc/template_preview', type='json', auth='user', methods=['POST'])
    def template_preview(self, doc_id, context=None, **kw):
        """以 user 提供的 context 渲染 doc.content_html，回傳完整 HTML 頁面供新分頁顯示。

        參數：
            doc_id：doc.document id
            context：{變數名稱: 值} dict（對應 content_html 中的 {{ var }}）

        回傳：
            { success: bool, html: str, warnings: list[str], error?: str }
        """
        try:
            doc = request.env['doc.document'].browse(int(doc_id))
            doc.check_access_rule('read')
        except Exception as e:
            return {'success': False, 'error': f'文件存取失敗：{e}'}

        content_html = doc.get_content_html() or '<p><em>（文件無內容）</em></p>'
        warnings = []
        ctx = context if isinstance(context, dict) else {}

        # 用 jinja2 sandbox 渲染（與 doc.render.mixin._render_template 同樣的 sandbox）
        rendered_body = content_html
        if ctx:
            try:
                from jinja2.sandbox import SandboxedEnvironment
                from jinja2 import StrictUndefined, UndefinedError
                env = SandboxedEnvironment(undefined=StrictUndefined)
                tpl = env.from_string(content_html)
                try:
                    rendered_body = tpl.render(**ctx, object=doc, user=request.env.user)
                except UndefinedError as ue:
                    warnings.append(f'缺少變數：{ue}（已用空白替代）')
                    env_lax = SandboxedEnvironment()
                    rendered_body = env_lax.from_string(content_html).render(
                        **ctx, object=doc, user=request.env.user
                    )
            except Exception as e:
                warnings.append(f'渲染警告：{e}')
                rendered_body = content_html

        # 包成完整 HTML 頁面（含列印樣式）
        page_format = doc.page_format or 'A4'
        doc_name = html_mod.escape(doc.name or '未命名文件')
        # html_mod.escape on warnings (avoid XSS via warning text)
        warnings_html = ''.join(
            f'<li>{html_mod.escape(w)}</li>' for w in warnings
        )
        warnings_block = (
            f'<aside class="preview-warnings" role="alert">'
            f'<strong>預覽提示</strong><ul>{warnings_html}</ul></aside>'
        ) if warnings else ''

        html = (
            '<!DOCTYPE html>'
            '<html lang="zh-Hant">'
            '<head>'
            '<meta charset="utf-8"/>'
            f'<title>預覽：{doc_name}</title>'
            '<style>'
            'body{font-family:"Microsoft JhengHei","PingFang TC",sans-serif;'
            'margin:0;padding:24px;background:#f3f4f6;color:#111827;}'
            '.preview-page{background:#fff;max-width:794px;margin:0 auto 16px;'
            'padding:48px 56px;box-shadow:0 1px 3px rgba(0,0,0,.1);'
            'border-radius:4px;line-height:1.6;}'
            '.preview-header{max-width:794px;margin:0 auto 16px;'
            'display:flex;justify-content:space-between;align-items:center;'
            'color:#4b5563;font-size:14px;}'
            '.preview-warnings{max-width:794px;margin:0 auto 16px;'
            'padding:12px 16px;background:#fef3c7;border-left:4px solid #f59e0b;'
            'border-radius:4px;color:#92400e;}'
            '.preview-warnings ul{margin:8px 0 0;padding-left:20px;}'
            '@media print{body{background:#fff;padding:0;}'
            '.preview-page{box-shadow:none;border:none;margin:0;}'
            '.preview-header,.preview-warnings{display:none;}}'
            '</style>'
            '</head>'
            '<body>'
            '<div class="preview-header">'
            f'<span>{doc_name}</span><span>{page_format} ｜ 預覽（非正式輸出）</span>'
            '</div>'
            f'{warnings_block}'
            f'<article class="preview-page">{rendered_body}</article>'
            '</body></html>'
        )
        return {'success': True, 'html': html, 'warnings': warnings}

    @http.route('/dobtor_doc/template_requests/list', type='json', auth='user', methods=['POST'])
    def template_requests_list(self, doc_id, **kw):
        """列出此文件範本的填寫請求清單。

        目前 doc.fill.request model 尚未實作；先回傳空陣列作為殼，
        前端 Sprint A 的「請求」分頁能正常顯示「目前沒有填寫請求」空態。
        未來接上 model 後改成 search_read。
        """
        try:
            doc = request.env['doc.document'].browse(int(doc_id))
            doc.check_access_rule('read')
        except Exception as e:
            return {'success': False, 'error': f'文件存取失敗：{e}', 'requests': []}

        FillRequest = request.env.get('doc.fill.request')
        if FillRequest is None:
            return {'success': True, 'requests': []}

        # model 已實作時走真實 search_read
        try:
            records = FillRequest.sudo().search_read(
                domain=[('doc_id', '=', doc.id)],
                fields=['id', 'recipient_name', 'state', 'sent_at', 'completed_at'],
                order='sent_at desc',
            )
            STATE_LABELS = {
                'draft': '草稿',
                'sent': '已送出',
                'opened': '已開啟',
                'completed': '已完成',
                'expired': '已過期',
            }
            for r in records:
                r['state_label'] = STATE_LABELS.get(r.get('state'), r.get('state') or '—')
                r['sent_at'] = (r.get('sent_at') and str(r['sent_at'])) or None
                r['completed_at'] = (r.get('completed_at') and str(r['completed_at'])) or None
            return {'success': True, 'requests': records}
        except Exception as e:
            _logger.warning('template_requests_list fallback: %s', e)
            return {'success': True, 'requests': []}
